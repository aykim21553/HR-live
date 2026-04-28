from __future__ import annotations

import base64
import io
import json
import re
import sys
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def extract_docx_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    lines = []
    for p in doc.paragraphs:
        text = " ".join(p.text.split())
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(" ".join(cell.text.split()) for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def markdownish_to_paragraphs(text: str) -> list[str]:
    cleaned = text.replace("\r\n", "\n")
    blocks = [block.strip() for block in re.split(r"\n\s*\n", cleaned) if block.strip()]
    return blocks or [cleaned.strip()]


def export_docx(title: str, body: str, output_path: Path) -> None:
    doc = Document()
    doc.add_heading(title, level=0)
    for block in markdownish_to_paragraphs(body):
        doc.add_paragraph(block)
    doc.save(str(output_path))


def export_pdf(title: str, body: str, output_path: Path) -> None:
    styles = getSampleStyleSheet()
    font_candidates = [
        Path("C:/Windows/Fonts/malgun.ttf"),
        Path("C:/Windows/Fonts/맑은 고딕.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    font_name = "Helvetica"
    for candidate in font_candidates:
        if candidate.exists():
            font_name = "CustomDocFont"
            pdfmetrics.registerFont(TTFont(font_name, str(candidate)))
            break
    styles["Title"].fontName = font_name
    styles["BodyText"].fontName = font_name
    story = [Paragraph(title, styles["Title"]), Spacer(1, 18)]
    for block in markdownish_to_paragraphs(body):
        safe = (
            block.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )
        story.append(Paragraph(safe, styles["BodyText"]))
        story.append(Spacer(1, 10))
    doc = SimpleDocTemplate(str(output_path), pagesize=A4)
    doc.build(story)


def main() -> int:
    if len(sys.argv) < 2:
        raise SystemExit("missing command")

    command = sys.argv[1]
    if command == "extract":
        path = Path(sys.argv[2])
        print(json.dumps({"text": extract_docx_text(path)}, ensure_ascii=False))
        return 0

    payload = json.load(sys.stdin)
    title = payload.get("title", "Contract Output")
    body = payload.get("body", "")
    output_path = Path(payload["outputPath"])
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if command == "export-docx":
        export_docx(title, body, output_path)
    elif command == "export-pdf":
        export_pdf(title, body, output_path)
    else:
        raise SystemExit(f"unsupported command: {command}")

    print(json.dumps({"ok": True, "path": str(output_path)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
