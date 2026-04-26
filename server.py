"""
server.py — HR JD RAG Scorer (Claude 버전)
  - LLM  : Anthropic Claude (claude-haiku-4-5-20251001 기본)
  - 임베딩: scikit-learn TF-IDF + 코사인 유사도 (로컬, 무료)
Run:  uvicorn server:app --reload --host 127.0.0.1 --port 8765
Env:  ANTHROPIC_API_KEY  (또는 .env 파일)
"""

from __future__ import annotations

import json
import os
import re
from pathlib import Path
from datetime import datetime, timedelta, timezone
from email.utils import parsedate_to_datetime
from urllib.parse import quote_plus
import xml.etree.ElementTree as ET

import httpx
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

# ── scikit-learn (TF-IDF 임베딩) ──────────────────────────────
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

# ── Anthropic ────────────────────────────────────────────────
import anthropic as _anthropic

load_dotenv()

ROOT = Path(__file__).resolve().parent

# ══════════════════════════════════════════════════════════════
# 학자금 데이터 경로 & 규정 상수
# ══════════════════════════════════════════════════════════════
_TUITION_DATA  = ROOT / "data"
_DB_FILE       = _TUITION_DATA / "payment_history_full.json"   # 7년 이력 DB
_APPS_FILE     = _TUITION_DATA / "applications.json"           # 신청 접수 store

TUITION_LIMITS         = {"유치원": 300_000,   "고등학교": 1_000_000,  "대학교": 3_200_000}
TUITION_OVERSEAS_LIMITS= {"유치원": 500_000,   "고등학교": 2_000_000,  "대학교": 6_000_000}
TUITION_MAX_TERMS      = {"유치원": 12,         "고등학교": 12,          "대학교": 8}   # 4년×2학기=8학기

# ── 회사 DB 로드 ──────────────────────────────────────────────
def _load_db() -> dict:
    if _DB_FILE.exists():
        with open(_DB_FILE, encoding="utf-8") as f:
            return json.load(f)
    return {"employees": {}, "payments": []}

_company_db: dict = _load_db()   # 서버 시작 시 1회 로드

# ── 신청 store 읽기/쓰기 ──────────────────────────────────────
import threading
_apps_lock = threading.Lock()

def _load_apps() -> list:
    if _APPS_FILE.exists():
        with open(_APPS_FILE, encoding="utf-8") as f:
            return json.load(f)
    return []

def _save_apps(apps: list) -> None:
    _TUITION_DATA.mkdir(parents=True, exist_ok=True)
    with open(_APPS_FILE, "w", encoding="utf-8") as f:
        json.dump(apps, f, ensure_ascii=False, indent=2, default=str)

# ── 학자금 DB 대조 엔진 ───────────────────────────────────────
def _name_match(a: str, b: str) -> bool:
    """이름 유사 비교 (완전일치 or 포함)"""
    if not a or not b:
        return False
    a, b = a.strip(), b.strip()
    return a == b or a in b or b in a

def cross_reference(employee_id: str, ocr: dict) -> dict:
    """
    OCR 추출 데이터 vs 회사 DB 대조.
    반환값: 규정 판정 + rule_checks + 이력 + flags + verdict
    """
    employees = _company_db.get("employees", {})
    payments  = _company_db.get("payments", [])

    child_name     = (ocr.get("child_name")    or "").strip()
    school_type    = (ocr.get("school_type")   or "").strip()
    school_name    = (ocr.get("school_name")   or "").strip()
    school_country = (ocr.get("school_country") or "KR").strip()
    benefit_year   = ocr.get("benefit_year") or 0
    benefit_term   = (ocr.get("benefit_term")  or "").strip()
    amount         = int(ocr.get("amount_total") or 0)
    payment_date   = (ocr.get("payment_date")  or "").strip()
    is_overseas    = school_country != "KR"

    r: dict = {
        # 직원 조회
        "employee_found":   False,
        "employee_name":    None,
        "employee_dept":    None,
        "employee_join_date": None,
        # 자녀 조회
        "child_found":      False,
        "child_registered": None,
        "child_name_match": False,
        # 학교 대조
        "school_type_match": False,
        "school_name_match": False,
        "school_name_registered": None,
        # 납부일
        "payment_date_ok":  True,
        "payment_date_days": 0,
        # 이력
        "history":          [],
        "duplicate":        False,
        "duplicate_records":[],
        "terms_used_ytd":   0,
        "terms_remaining":  0,
        "amount_used_ytd":  0,
        # 한도
        "limit_per_term":   0,
        "amount_payable":   0,
        "over_limit":       False,
        "is_overseas":      is_overseas,
        # 판정
        "flags":            [],
        "rule_checks":      {},
        "verdict":          "수기검토필요",
        "verdict_reason":   "",
    }

    # ── 1. 직원 확인 ─────────────────────────────────────────
    emp = employees.get(employee_id)
    if emp:
        r["employee_found"]    = True
        r["employee_name"]     = emp.get("name")
        r["employee_dept"]     = emp.get("dept")
        r["employee_join_date"]= emp.get("join_date")

    # ── 2. 자녀 등록 확인 ────────────────────────────────────
    matched_child = None
    if emp:
        for ch in emp.get("children", []):
            if _name_match(child_name, ch.get("name", "")):
                matched_child = ch
                r["child_found"]   = True
                r["child_registered"] = ch
                r["child_name_match"] = ch.get("name") == child_name  # 완전일치 여부
                break

    # ── 3. 학교 대조 ─────────────────────────────────────────
    if matched_child:
        r["school_name_registered"] = matched_child.get("school_name")
        r["school_type_match"]  = matched_child.get("school_type") == school_type
        reg_sn = (matched_child.get("school_name") or "").strip().lower()
        ocr_sn = school_name.lower()
        r["school_name_match"] = (reg_sn == ocr_sn) or (reg_sn and ocr_sn and (reg_sn in ocr_sn or ocr_sn in reg_sn))

    # ── 4. 납부일 기한 (90일 이내) ───────────────────────────
    if payment_date:
        try:
            from datetime import date as _date
            pd  = _date.fromisoformat(payment_date)
            today = _date.today()
            diff  = (today - pd).days
            r["payment_date_ok"]   = 0 <= diff <= 90
            r["payment_date_days"] = diff
        except Exception:
            pass

    # ── 5. 이력 조회 ─────────────────────────────────────────
    child_id = matched_child.get("id") if matched_child else None
    child_hist = [p for p in payments if p.get("child_id") == child_id] if child_id else []
    r["history"] = child_hist

    # ── 6. 중복 수혜 ─────────────────────────────────────────
    dups = [p for p in child_hist
            if p.get("year") == benefit_year
            and p.get("benefit_term") == benefit_term
            and p.get("status") == "approved"]
    r["duplicate"]         = len(dups) > 0
    r["duplicate_records"] = dups

    # ── 7. 연간 누적 ─────────────────────────────────────────
    ytd = [p for p in child_hist if p.get("year") == benefit_year and p.get("status") == "approved"]
    r["terms_used_ytd"]  = len(ytd)
    r["amount_used_ytd"] = sum(p.get("amount_paid", 0) for p in ytd)
    max_t = TUITION_MAX_TERMS.get(school_type, 12)
    r["terms_remaining"] = max(0, max_t - r["terms_used_ytd"])

    # ── 8. 한도 ──────────────────────────────────────────────
    limit = (TUITION_OVERSEAS_LIMITS if is_overseas else TUITION_LIMITS).get(school_type, 0)
    r["limit_per_term"]  = limit
    r["amount_payable"]  = min(amount, limit)
    r["over_limit"]      = amount > limit and limit > 0

    # ── 9. Flags ─────────────────────────────────────────────
    flags = []
    if not r["employee_found"]:     flags.append("직원미등록")
    if not r["child_found"]:        flags.append("자녀정보미등록")
    elif not r["child_name_match"]: flags.append("자녀명_부분일치")
    if not r["school_type_match"] and matched_child:
        flags.append(f"학교급불일치")
    if not r["school_name_match"] and matched_child:
        flags.append("학교명불일치")
    if not r["payment_date_ok"]:    flags.append("납부일기한초과")
    if r["duplicate"]:              flags.append(f"중복수혜_{benefit_year}-{benefit_term}")
    if r["terms_remaining"] == 0:   flags.append("연간한도소진")
    if r["over_limit"]:             flags.append(f"한도초과_{amount - limit:,}원")
    if is_overseas:                 flags.append("해외학교_추가서류필요")
    r["flags"] = flags

    # ── 10. Rule checks ──────────────────────────────────────
    emp_dept = r["employee_dept"] or ""
    rc = {
        "직원 재직 확인":  {"passed": r["employee_found"],
            "reason": f"{employee_id} · {r['employee_name'] or '미등록'} ({emp_dept})"},
        "자녀 등록 확인":  {"passed": r["child_found"],
            "reason": f"신청: '{child_name}' → DB: '{matched_child.get('name','—') if matched_child else '미등록'}' ({'완전일치' if r['child_name_match'] else '유사일치' if r['child_found'] else '미등록'})"},
        "학교급 일치":     {"passed": r["school_type_match"] if matched_child else True,
            "reason": f"DB등록: {matched_child.get('school_type','—') if matched_child else '—'} / 신청: {school_type}"},
        "학교명 일치":     {"passed": r["school_name_match"] if matched_child else True,
            "reason": f"DB등록: {r['school_name_registered'] or '—'} / 신청서: {school_name}"},
        "납부일 기한":     {"passed": r["payment_date_ok"],
            "reason": f"납부일 {payment_date} ({r['payment_date_days']}일 경과, 기준 90일)"},
        "중복 수혜":       {"passed": not r["duplicate"],
            "reason": "동일 기간 지급 이력 없음" if not r["duplicate"] else f"{benefit_year}-{benefit_term} 기 이미 지급 ({len(dups)}건)"},
        "연간 한도":       {"passed": r["terms_remaining"] > 0,
            "reason": f"잔여 {r['terms_remaining']}회 / 올해 {r['terms_used_ytd']}회 사용 (최대 {max_t}회/년)"},
        "회당 지급 한도":  {"passed": not r["over_limit"],
            "reason": f"신청 {amount:,}원 {'≤' if not r['over_limit'] else '>'} 한도 {limit:,}원" +
                     (f" → {r['amount_payable']:,}원으로 조정" if r["over_limit"] else "")},
    }
    if is_overseas:
        rc["해외학교 추가서류"] = {"passed": None,
            "reason": f"국가: {school_country} — 재학증명서·납입영수증 원본·가족관계증명서 제출 필요"}
    r["rule_checks"] = rc

    # ── 11. Verdict ──────────────────────────────────────────
    hard_fail = [k for k,v in rc.items() if v["passed"] is False
                 and k in ("직원 재직 확인","자녀 등록 확인","연간 한도","중복 수혜")]
    soft_fail = [k for k,v in rc.items() if v["passed"] is False and k not in hard_fail]
    needs_rev = [k for k,v in rc.items() if v["passed"] is None]

    if not r["employee_found"] or not r["child_found"]:
        r["verdict"] = "반려"
        r["verdict_reason"] = "직원 또는 자녀 정보가 회사 DB에 등록되지 않아 지급이 불가합니다."
    elif "연간 한도" in hard_fail:
        r["verdict"] = "반려"
        r["verdict_reason"] = f"연간 지원 한도 소진 ({max_t}회 모두 사용)으로 추가 지급이 불가합니다."
    elif "중복 수혜" in hard_fail:
        r["verdict"] = "수기검토필요"
        r["verdict_reason"] = f"{benefit_year}년 {benefit_term} 동일 기간 지급 이력이 존재합니다. 중복 수혜 여부를 확인하세요."
    elif needs_rev:
        r["verdict"] = "수기검토필요"
        r["verdict_reason"] = "해외학교 추가 서류 원본 제출 후 수기 검토가 필요합니다."
    elif "학교명 일치" in soft_fail:
        r["verdict"] = "보완필요"
        r["verdict_reason"] = f"신청서 학교명({school_name})이 DB 등록 학교명({r['school_name_registered']})과 다릅니다. 정확한 서류를 재제출하세요."
    elif soft_fail:
        r["verdict"] = "자동승인후보"
        r["verdict_reason"] = f"한도 조정 후 {r['amount_payable']:,}원 지급 예정. " + "; ".join(flags)
    else:
        r["verdict"] = "자동승인후보"
        r["verdict_reason"] = f"모든 항목 검증 통과. 지급 예정: {r['amount_payable']:,}원"

    return r
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
CHAT_MODEL        = os.getenv("CLAUDE_MODEL", "claude-haiku-4-5-20251001")
# OCR은 정확도 우선 -> Sonnet 4.6 사용 (Haiku는 표·세로 라벨 해석에서 실패율 높음)
OCR_MODEL         = os.getenv("OCR_MODEL",    "claude-sonnet-4-6")
CHUNK_SIZE        = int(os.getenv("RAG_CHUNK_SIZE", "720"))
CHUNK_OVERLAP     = int(os.getenv("RAG_CHUNK_OVERLAP", "100"))
TOP_K             = int(os.getenv("RAG_TOP_K", "5"))
RESUME_LLM_CHARS  = int(os.getenv("RESUME_LLM_CHARS", "14000"))

app = FastAPI(title="HR JD RAG Scorer — Claude Edition")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*", "null"],   # "null" = file:// 로 열었을 때 브라우저가 보내는 origin
    allow_credentials=False,        # wildcard origin 과 credentials 동시 사용 불가
    allow_methods=["*"],
    allow_headers=["*"],
)

# ══════════════════════════════════════════════════════════════
# 보안 미들웨어 — Rate Limiting
# ══════════════════════════════════════════════════════════════
from collections import defaultdict
import time as _time

_rate_store: dict = defaultdict(list)
RATE_LIMIT_CALLS = 30   # 분당 최대 요청 수
RATE_LIMIT_WINDOW = 60  # 초

def _check_rate_limit(client_ip: str) -> bool:
    """True = 허용, False = 차단"""
    now = _time.time()
    window_start = now - RATE_LIMIT_WINDOW
    calls = [t for t in _rate_store[client_ip] if t > window_start]
    if len(calls) >= RATE_LIMIT_CALLS:
        return False
    calls.append(now)
    _rate_store[client_ip] = calls
    return True

from fastapi import Request as _Request

@app.middleware("http")
async def rate_limit_middleware(request: _Request, call_next):
    client_ip = request.client.host if request.client else "unknown"
    # API 엔드포인트에만 적용
    if request.url.path.startswith("/api/"):
        if not _check_rate_limit(client_ip):
            from fastapi.responses import JSONResponse
            return JSONResponse(
                status_code=429,
                content={"detail": "요청 한도 초과. 1분 후 다시 시도하세요."}
            )
    response = await call_next(request)
    # 보안 헤더 추가
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "SAMEORIGIN"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    return response




def get_client() -> _anthropic.Anthropic:
    if not ANTHROPIC_API_KEY:
        raise HTTPException(
            status_code=503,
            detail="서버에 ANTHROPIC_API_KEY 환경 변수가 설정되어 있지 않습니다. .env 파일을 확인하세요.",
        )
    return _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)


# ── HTML 파싱 ─────────────────────────────────────────────────

def strip_html(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "template"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    lines = [re.sub(r"\s+", " ", ln).strip() for ln in text.splitlines()]
    return "\n".join(ln for ln in lines if ln)


def _safe_pub_dt(v: str | None) -> datetime | None:
    if not v:
        return None
    try:
        dt = parsedate_to_datetime(v)
        if dt.tzinfo is None:
            return dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(timezone.utc)
    except Exception:
        return None


async def _fetch_google_news_rss(query: str, days: int = 7, limit: int = 25) -> list[dict]:
    q = quote_plus(f"{query} when:{max(1, days)}d")
    url = f"https://news.google.com/rss/search?q={q}&hl=ko&gl=KR&ceid=KR:ko"
    try:
        async with httpx.AsyncClient(timeout=15.0, follow_redirects=True) as hx:
            r = await hx.get(url, headers={"User-Agent": "Mozilla/5.0 (HR Newsroom Bot)"})
            r.raise_for_status()
            xml_text = r.text
    except Exception:
        return []

    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        return []

    items: list[dict] = []
    for item in root.findall("./channel/item"):
        title = (item.findtext("title") or "").strip()
        link = (item.findtext("link") or "").strip()
        pub = _safe_pub_dt(item.findtext("pubDate"))
        source = ""
        src_el = item.find("source")
        if src_el is not None and src_el.text:
            source = src_el.text.strip()
        if not title or not link:
            continue
        items.append(
            {
                "title": title,
                "url": link,
                "source": source or "Google News",
                "published_at": pub.isoformat() if pub else "",
            }
        )
        if len(items) >= limit:
            break
    return items


def _dedupe_items(items: list[dict]) -> list[dict]:
    out: list[dict] = []
    seen = set()
    for it in items:
        key = (it.get("title", "")[:80], it.get("url", ""))
        if key in seen:
            continue
        seen.add(key)
        out.append(it)
    return out


# ── 텍스트 청킹 ───────────────────────────────────────────────

def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    t = re.sub(r"\s+", " ", text).strip()
    if not t:
        return []
    if len(t) <= size:
        return [t]
    chunks: list[str] = []
    start = 0
    while start < len(t):
        end = min(start + size, len(t))
        chunks.append(t[start:end].strip())
        if end >= len(t):
            break
        start = max(0, end - overlap)
    return [c for c in chunks if c]


# ── TF-IDF 기반 청크 선택 ─────────────────────────────────────

def pick_top_chunks_tfidf(
    chunks: list[str],
    resume_text: str,
    k: int,
) -> list[tuple[int, str, float]]:
    """
    TF-IDF 벡터화 후 코사인 유사도로 상위 k개 청크 반환.
    scikit-learn 사용 — API 비용 없음, 완전 로컬 실행.
    """
    if not chunks:
        return []

    corpus = chunks + [resume_text]
    vectorizer = TfidfVectorizer(
        analyzer="char_wb",   # n-gram 기반 → 한국어·영어 혼용에 강함
        ngram_range=(2, 4),
        max_features=20000,
        sublinear_tf=True,
    )
    try:
        tfidf_matrix = vectorizer.fit_transform(corpus)
    except ValueError:
        # 빈 vocabulary 등 엣지 케이스
        return [(i, chunks[i], 0.0) for i in range(min(k, len(chunks)))]

    chunk_vecs  = tfidf_matrix[: len(chunks)]
    resume_vec  = tfidf_matrix[len(chunks)]

    sims = cosine_similarity(resume_vec, chunk_vecs).flatten()
    top_idx = np.argsort(sims)[::-1][: k]

    return [(int(i), chunks[i], float(sims[i])) for i in top_idx]


# ── Claude LLM 스코어링 ───────────────────────────────────────

def llm_score_one(
    client: _anthropic.Anthropic,
    jd_excerpts: list[str],
    resume_text: str,
    filename: str,
) -> dict:
    system_prompt = (
        "당신은 한화투자증권 채용 담당자를 보조하는 전문 HR 애널리스트입니다. "
        "아래 [직무 레퍼런스 발췌]는 JD/공고에서 TF-IDF RAG로 뽑힌 조각입니다. "
        "이력서와의 직무 밀접도를 0~100 정수로 평가하고, 한국어로 JSON만 출력하세요. "
        "JSON 외 다른 텍스트는 절대 출력하지 마세요."
    )
    ref_block = "\n\n---\n\n".join(
        f"[발췌 {i+1}]\n{s}" for i, s in enumerate(jd_excerpts)
    )
    user_prompt = (
        f"지원 파일명: {filename}\n\n"
        f"[직무 레퍼런스 발췌]\n{ref_block}\n\n"
        f"[이력서 본문]\n{resume_text[:RESUME_LLM_CHARS]}\n\n"
        "출력 JSON 키: score(0~100 정수), verdict(한 줄 판단), summary(3문장 이내), "
        "strengths(문자열 배열 최대 5), gaps(문자열 배열 최대 5), "
        "matched_reference_points(레퍼런스와 맞닿은 점, 문자열 배열 최대 5)."
    )

    try:
        message = client.messages.create(
            model=CHAT_MODEL,
            max_tokens=1024,
            messages=[{"role": "user", "content": user_prompt}],
            system=system_prompt,
        )
        raw = message.content[0].text if message.content else "{}"
    except _anthropic.AuthenticationError:
        raise HTTPException(status_code=401,
            detail="ANTHROPIC_API_KEY 가 유효하지 않습니다. .env 파일의 키 전체를 확인하세요.")
    except _anthropic.RateLimitError:
        raise HTTPException(status_code=429,
            detail="Anthropic API 요청 한도를 초과했습니다. 잠시 후 다시 시도하세요.")
    except _anthropic.APIConnectionError as e:
        raise HTTPException(status_code=502,
            detail=f"Anthropic 서버 연결 실패: {e}")
    except _anthropic.APIStatusError as e:
        raise HTTPException(status_code=e.status_code,
            detail=f"Anthropic API 오류 ({e.status_code}): {e.message}")

    # JSON 블록 추출 (```json ... ``` 등 감쌌을 경우 대비)
    json_match = re.search(r"\{[\s\S]*\}", raw)
    raw_json = json_match.group(0) if json_match else raw

    try:
        return json.loads(raw_json)
    except json.JSONDecodeError:
        return {
            "score": 0,
            "verdict": "JSON 파싱 실패",
            "summary": raw[:500],
            "strengths": [],
            "gaps": [],
            "matched_reference_points": [],
        }


# ══════════════════════════════════════════════════════════════
# 학자금 신청 API (applications.json 기반)
# ══════════════════════════════════════════════════════════════

class TuitionSubmitBody(BaseModel):
    employee_id:  str
    ocr:          dict          # OCR 추출 결과 전체
    file_data:    str = ""      # base64 인코딩된 원본 첨부 파일
    filename:     str = ""      # 첨부 파일 원본명 (확장자로 미디어 타입 추론)


@app.post("/api/tuition/submit")
def tuition_submit(body: TuitionSubmitBody):
    """직원이 신청 제출 → DB 대조 자동 실행 → applications.json 저장"""
    import uuid, datetime as _dt
    app_id  = "APP-" + datetime.now().strftime("%Y%m%d") + "-" + str(uuid.uuid4())[:6].upper()
    cross   = cross_reference(body.employee_id, body.ocr)
    record  = {
        "id":           app_id,
        "employee_id":  body.employee_id,
        "employee_name": cross.get("employee_name") or body.employee_id,
        "submitted_at": datetime.now().isoformat(),
        "status":       "pending",
        "ocr":          body.ocr,
        "cross_ref":    cross,
        "admin_action": None,
        "admin_comment":"",
        "draft_text":   None,
        "filename":     body.filename,
        "file_data":    body.file_data,   # base64 저장 (첨부파일 뷰어용)
    }
    with _apps_lock:
        apps = _load_apps()
        apps.insert(0, record)   # 최신순
        _save_apps(apps)
    return {"id": app_id, "verdict": cross["verdict"], "flags": cross["flags"]}


@app.get("/api/tuition/admin/applications")
def list_applications():
    """관리자: 전체 신청 목록 (최신순)"""
    with _apps_lock:
        return _load_apps()


@app.get("/api/tuition/admin/applications/{app_id}")
def get_application(app_id: str):
    """관리자: 특정 신청 상세 (cross_ref 포함)"""
    with _apps_lock:
        apps = _load_apps()
    for a in apps:
        if a["id"] == app_id:
            return a
    raise HTTPException(status_code=404, detail=f"신청 {app_id}를 찾을 수 없습니다.")


class AdminActionBody(BaseModel):
    action:  str   # approved | partial | rejected | supplement
    comment: str = ""
    amount_final: int = 0   # 일부 승인 시 실지급액


@app.patch("/api/tuition/admin/applications/{app_id}/action")
def admin_action(app_id: str, body: AdminActionBody):
    """관리자 처리 결과 저장"""
    with _apps_lock:
        apps = _load_apps()
        for a in apps:
            if a["id"] == app_id:
                a["status"]        = body.action
                a["admin_action"]  = body.action
                a["admin_comment"] = body.comment
                a["processed_at"]  = datetime.now().isoformat()
                if body.amount_final:
                    a["amount_final"] = body.amount_final
                _save_apps(apps)
                return {"ok": True, "id": app_id, "status": body.action}
    raise HTTPException(status_code=404, detail=f"신청 {app_id}를 찾을 수 없습니다.")


@app.post("/api/tuition/admin/applications/{app_id}/draft")
def generate_draft_api(app_id: str):
    """기안문 생성 (Claude LLM) 및 저장"""
    with _apps_lock:
        apps = _load_apps()
    app_rec = next((a for a in apps if a["id"] == app_id), None)
    if not app_rec:
        raise HTTPException(status_code=404, detail="신청 없음")

    ocr  = app_rec.get("ocr",      {})
    cr   = app_rec.get("cross_ref", {})
    emp  = cr.get("employee_name") or app_rec.get("employee_id", "")

    # LLM 기안문 생성 시도 (없으면 템플릿)
    draft = None
    if ANTHROPIC_API_KEY:
        try:
            client = _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
            prompt = f"""다음 학자금 지원 신청 정보를 바탕으로 회사 내부 기안문을 작성하세요.
중요: 마크다운(###, **, *, |--|, ``` 등) 일절 사용 금지. 순수 텍스트만 출력하세요.
공식적이고 간결한 한국어 문체를 사용하고, 아래 형식을 따르세요.

신청정보:
- 신청번호: {app_id}
- 신청자: {emp} ({app_rec.get('employee_id','')})
- 수혜자(자녀명): {ocr.get('child_name','—')}
- 학교명: {ocr.get('school_name','—')} ({ocr.get('school_type','—')})
- 납부일: {ocr.get('payment_date','—')}
- 수혜기간: {ocr.get('benefit_year','—')}년 {ocr.get('benefit_term','—')}
- 납부금액: {(ocr.get('amount_total') or 0):,}원
- 지급예정금액: {(cr.get('amount_payable') or 0):,}원 (한도 적용)
- 판정결과: {cr.get('verdict','—')}
- DB대조결과: {cr.get('verdict_reason','—')}
- 특이사항: {', '.join(cr.get('flags',[]) or ['없음'])}

형식 규칙 (반드시 준수):
- 수신: 신청자 이름 ({emp}) 으로 고정
- 참조: 항목 자체를 생략 (출력하지 말 것)
- 작성부서: 인재관리팀 으로 고정
- 본문 구성: 신청개요 → 납부내역 → 지원금액 → 처리의견
- 작성일 포함"""
            msg = client.messages.create(model=CHAT_MODEL, max_tokens=1500,
                messages=[{"role":"user","content":prompt}])
            draft = msg.content[0].text if msg.content else None
        except Exception:
            pass

    if not draft:
        draft = _build_draft_template(app_id, emp, app_rec.get("employee_id",""), ocr, cr)

    # 저장
    with _apps_lock:
        apps2 = _load_apps()
        for a in apps2:
            if a["id"] == app_id:
                a["draft_text"] = draft
                break
        _save_apps(apps2)
    return {"draft_text": draft}


@app.get("/api/tuition/admin/applications/{app_id}/file")
def get_application_file(app_id: str):
    """첨부 파일 서빙 (base64 → binary Response)"""
    import base64 as _b64
    from fastapi.responses import Response as _Resp
    with _apps_lock:
        apps = _load_apps()
    rec = next((a for a in apps if a["id"] == app_id), None)
    if not rec:
        raise HTTPException(status_code=404, detail="신청 없음")
    fd = rec.get("file_data", "")
    fn = rec.get("filename", "")
    if not fd:
        raise HTTPException(status_code=404, detail="첨부 파일 없음")
    if "," in fd:
        fd = fd.split(",", 1)[1]
    try:
        raw = _b64.b64decode(fd)
    except Exception:
        raise HTTPException(status_code=422, detail="파일 디코딩 실패")
    fn_lower = fn.lower()
    if fn_lower.endswith(".pdf"):
        mt = "application/pdf"
    elif fn_lower.endswith(".png"):
        mt = "image/png"
    elif fn_lower.endswith(".gif"):
        mt = "image/gif"
    elif fn_lower.endswith(".webp"):
        mt = "image/webp"
    else:
        mt = "image/jpeg"
    return _Resp(content=raw, media_type=mt,
                 headers={"Content-Disposition": f'inline; filename="{fn or "attachment"}"'})


class BatchDraftBody(BaseModel):
    app_ids: list  # 처리할 app id 목록 (비어있으면 자동승인후보 전체)


@app.post("/api/tuition/admin/batch-draft")
def batch_draft(body: BatchDraftBody):
    """일괄 기안문 생성 — 자동승인후보 대상"""
    with _apps_lock:
        apps = _load_apps()

    if body.app_ids:
        targets = [a for a in apps if a["id"] in body.app_ids]
    else:
        targets = [a for a in apps
                   if (a.get("cross_ref") or {}).get("verdict") == "자동승인후보"]

    results = []
    for app_rec in targets:
        ocr = app_rec.get("ocr", {})
        cr  = app_rec.get("cross_ref", {})
        emp = cr.get("employee_name") or app_rec.get("employee_id", "")
        draft = _build_draft_template(app_rec["id"], emp, app_rec.get("employee_id",""), ocr, cr)
        results.append({
            "id": app_rec["id"],
            "employee_name": emp,
            "child_name": ocr.get("child_name","—"),
            "verdict": cr.get("verdict","—"),
            "draft_text": draft,
        })

    with _apps_lock:
        apps2 = _load_apps()
        id_map = {r["id"]: r["draft_text"] for r in results}
        for a in apps2:
            if a["id"] in id_map:
                a["draft_text"] = id_map[a["id"]]
        _save_apps(apps2)

    return {"count": len(results), "drafts": results}


def _build_draft_template(app_id, emp_name, emp_id, ocr, cr):
    from datetime import date
    today = date.today().strftime("%Y년 %m월 %d일")
    is_ov = cr.get("is_overseas", False)
    flags = cr.get("flags") or []
    flag_note = "\n   ※ 특이사항: " + ", ".join(flags) if flags else ""
    ov_note = f"\n\n5. 해외학교 특이사항\n   ▸ 국가코드  : {ocr.get('school_country','—')}\n   ▸ 적용한도  : 해외 특례 한도 적용\n   ▸ 추가서류  : 재학증명서·납입영수증 원본 확인 필요" if is_ov else ""
    return f"""                학 자 금 지 급 신 청 기 안

─────────────────────────────────────────────────────
수    신 : {emp_name} ({emp_id})
기 안 일 : {today}
기 안 자 : 인재관리팀
결 재 선 : 팀장 → 본부장 (학자금 지원 규정 제12조)
─────────────────────────────────────────────────────

1. 신청 개요
   ▸ 신청번호 : {app_id}
   ▸ 수혜자   : {ocr.get('child_name','—')} ({ocr.get('school_type','—')})
   ▸ 학교명   : {ocr.get('school_name','—')}{'  [해외학교]' if is_ov else ''}
   ▸ 수혜기간 : {ocr.get('benefit_year','—')}년 {ocr.get('benefit_term','—')}
   ▸ 서류종류 : {ocr.get('doc_type','—')}

2. 납부 내역
   ▸ 납부일자     : {ocr.get('payment_date','—')}
   ▸ 납부 합계    : {(ocr.get('amount_total') or 0):,}원
   ▸ 지급 예정    : {(cr.get('amount_payable') or 0):,}원  (규정 한도 적용)

3. DB 대조 결과
   ▸ 판정결과   : {cr.get('verdict','—')}
   ▸ 판정근거   : {cr.get('verdict_reason','—')}{flag_note}
   ▸ 연간 누적  : {cr.get('terms_used_ytd',0)}회 사용 / 잔여 {cr.get('terms_remaining',0)}회

4. 처리 의견
   위 학자금 지원 신청에 대해 회사 DB 대조 및 자동 규정 검토를
   완료하였기에 결재를 요청합니다.{ov_note}

─────────────────────────────────────────────────────
담당자: ______________  팀장: ______________  본부장: ______________
─────────────────────────────────────────────────────
작성부서: 인재관리팀"""


# ══════════════════════════════════════════════════════════════
# Ellis in Wonderland — 팀장 피드백 분석 API
# ══════════════════════════════════════════════════════════════

class EllisMemberFeedback(BaseModel):
    name:       str = ""
    grade:      str = ""          # S/A/B/C
    job:        str = ""          # 본업 피드백
    innovation: str = ""          # 혁신 피드백
    growth:     str = ""          # 성장 피드백

class EllisAnalyzeBody(BaseModel):
    manager_name: str = "팀장"
    members: list[EllisMemberFeedback]


@app.post("/api/ellis/analyze")
async def ellis_analyze(body: EllisAnalyzeBody):
    """
    팀장 피드백 5개 축 분석:
    객관성 / 공정성 / 균형성 / 일관성 / 피드백 문화 적합성
    """
    if not body.members:
        raise HTTPException(status_code=400, detail="팀원 피드백 데이터가 없습니다.")

    client = get_client()

    # 멤버 데이터 직렬화 — 피드백 텍스트는 토큰 절약을 위해 200자로 트런케이션
    def trim(text: str, limit: int = 200) -> str:
        text = (text or "").strip()
        return text[:limit] + ("…" if len(text) > limit else "")

    members_text = ""
    for i, m in enumerate(body.members):
        members_text += (
            f"\n[팀원{i+1}] {m.name or f'팀원{i+1}'}\n"
            f"  본업: {trim(m.job)}\n"
            f"  혁신: {trim(m.innovation)}\n"
            f"  성장: {trim(m.growth)}\n"
        )

    system_msg = (
        "당신은 HR 피드백 문화 전문가입니다. "
        "사용자의 요청에 대해 반드시 유효한 JSON 객체만 출력하세요. "
        "```json 같은 마크다운 코드블록, 설명 텍스트, 주석은 절대 포함하지 마세요. "
        "응답의 첫 글자는 반드시 { 여야 합니다."
    )

    prompt = f"""팀원 {len(body.members)}명의 팀장 피드백을 아래 5개 축으로 분석하고 JSON만 반환하세요.

분석 축:
1. objectivity(객관성): 사실·근거 기반 여부, 추상/성격판단 표현 탐지
2. fairness(공정성): 팀원 간 서술 강도·밀도 일관성
3. balance(균형성): 강점·개선점 균형, 코칭 문장 포함 여부
4. consistency(일관성): 복붙 의심, 분량 편차, 동일 행동 다른 해석
5. culture_fit(문화적합성): 비난·낙인형 표현 탐지, 행동중심·성장지향 여부

팀원 피드백:
{members_text}

반환 JSON (실제 분석 결과로 채우세요. 문자열은 간결하게 1~2줄 이내):
{{
  "summary":{{"overall_score":75,"grade":"양호","headline":"한 줄 요약","top_strength":"잘된 점","top_risk":"주의할 점"}},
  "dimensions":{{
    "objectivity":{{"score":75,"label":"객관성","finding":"소견 1~2줄","risk_expressions":["표현1"],"good_examples":["예시1"]}},
    "fairness":{{"score":75,"label":"공정성","finding":"소견 1~2줄","outlier_members":[]}},
    "balance":{{"score":75,"label":"균형성","finding":"소견 1~2줄","coaching_ratio":"60%"}},
    "consistency":{{"score":75,"label":"일관성","finding":"소견 1~2줄","copy_paste_risk":false,"length_variance":"보통"}},
    "culture_fit":{{"score":75,"label":"문화 적합성","finding":"소견 1~2줄","high_risk_expressions":[],"coaching_suggestions":["제안1"]}}
  }},
  "member_analysis":[
    {{"name":"팀원A","scores":{{"objectivity":75,"balance":70,"culture_fit":80}},"flags":[],"strength":"강점 한 줄","improvement":"개선 한 줄"}}
  ],
  "action_items":[
    {{"priority":"높음","item":"권고사항 한 줄"}}
  ]
}}"""

    try:
        msg = client.messages.create(
            model=CHAT_MODEL,
            max_tokens=8000,
            system=system_msg,
            messages=[{"role": "user", "content": prompt}]
        )
        raw = (msg.content[0].text if msg.content else "").strip()

        import ast

        def try_parse(text: str):
            """JSON 또는 Python dict 문법을 모두 허용하는 안전한 파서"""
            text = text.strip()
            # 1) 표준 JSON
            try:
                return json.loads(text)
            except (json.JSONDecodeError, ValueError):
                pass
            # 2) Python dict (single quote, True/False/None 등)
            try:
                val = ast.literal_eval(text)
                if isinstance(val, dict):
                    # dict → JSON 왕복으로 정규화
                    return json.loads(json.dumps(val, ensure_ascii=False))
            except Exception:
                pass
            return None

        result = None

        # 코드블록 제거
        cleaned = re.sub(r"```(?:json|python)?\s*", "", raw)
        cleaned = re.sub(r"```", "", cleaned).strip()

        # ① 전체 텍스트 파싱
        result = try_parse(cleaned)

        # ② 첫 { ~ 마지막 } 추출
        if result is None:
            s = cleaned.find("{")
            e = cleaned.rfind("}") + 1
            if s >= 0 and e > s:
                result = try_parse(cleaned[s:e])

        # ③ 정규식으로 가장 큰 블록
        if result is None:
            for candidate in sorted(re.findall(r"\{[\s\S]+\}", cleaned), key=len, reverse=True):
                result = try_parse(candidate)
                if result:
                    break

        if result is None:
            preview = raw[:400].replace('"', "'")
            raise HTTPException(
                status_code=422,
                detail=f"JSON 파싱 실패. 모델 응답 미리보기: {preview}"
            )

        return {"ok": True, "manager": body.manager_name, "analysis": result}

    except _anthropic.AuthenticationError:
        raise HTTPException(status_code=401, detail="API 키가 유효하지 않습니다.")
    except _anthropic.RateLimitError:
        raise HTTPException(status_code=429, detail="API 요청 한도 초과. 잠시 후 재시도하세요.")
    except _anthropic.APIStatusError as e:
        raise HTTPException(status_code=e.status_code, detail=f"API 오류 ({e.status_code}): {e.message}")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"서버 오류: {str(e)}")


# ── Ellis: 파일 업로드 → 피드백 자동 파싱 ────────────────────

def _extract_text_from_pdf(data: bytes) -> str:
    """pdfplumber로 PDF 텍스트 추출"""
    import io
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages)
    except Exception as e:
        return f"[PDF 추출 오류: {e}]"


def _extract_text_from_docx(data: bytes) -> str:
    """python-docx로 Word 텍스트 추출"""
    import io
    try:
        from docx import Document as D
        doc = D(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        return f"[DOCX 추출 오류: {e}]"


def _extract_text_from_xlsx(data: bytes) -> str:
    """openpyxl로 Excel 내용 추출 — 구조화 시도 후 텍스트 반환"""
    import io
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
        lines = []
        for ws in wb.worksheets:
            lines.append(f"[시트: {ws.title}]")
            for row in ws.iter_rows(values_only=True):
                vals = [str(c) if c is not None else "" for c in row]
                if any(v.strip() for v in vals):
                    lines.append("\t".join(vals))
        return "\n".join(lines)
    except Exception as e:
        return f"[Excel 추출 오류: {e}]"


def _parse_feedback_with_llm(client, raw_text: str, filename: str) -> list:
    """Claude로 원본 텍스트에서 팀원별 피드백 구조 추출"""
    prompt = f"""다음은 '{filename}' 파일에서 추출한 텍스트입니다.
이 내용에서 팀원별 피드백 정보를 추출하여 JSON 배열로 반환하세요.
JSON 외 다른 텍스트는 출력하지 마세요.

각 팀원별 필드:
- name: 팀원 이름 (성+이름)
- grade: 등급 (S/A/B/C 중 하나, 없으면 빈 문자열)
- job: 본업(담당업무/성과) 관련 피드백 텍스트
- innovation: 혁신/도전/창의 관련 피드백 텍스트
- growth: 성장/역량개발/학습 관련 피드백 텍스트

※ 본업/혁신/성장 구분이 없는 경우, 전체 피드백을 job 필드에 담고 나머지는 빈 문자열.
※ 팀원이 식별되지 않으면 빈 배열 [] 반환.

출력 형식: [{{"name":"..","grade":"..","job":"..","innovation":"..","growth":".."}}, ...]

원본 텍스트:
{raw_text[:8000]}"""

    try:
        msg = client.messages.create(
            model=CHAT_MODEL, max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )
        raw = msg.content[0].text if msg.content else "[]"
        m = re.search(r"\[[\s\S]*\]", raw)
        return json.loads(m.group(0)) if m else []
    except Exception:
        return []


@app.post("/api/ellis/parse-file")
async def ellis_parse_file(file: UploadFile = File(...)):
    """
    Excel / Word / PDF 업로드 → 팀원별 피드백 자동 파싱
    반환: { members: [{name, grade, job, innovation, growth}], source_text_preview }
    """
    data = await file.read()
    fname = (file.filename or "").lower()

    # 1. 파일 형식별 텍스트 추출
    if fname.endswith((".xlsx", ".xls", ".xlsm")):
        raw = _extract_text_from_xlsx(data)
        file_type = "Excel"
    elif fname.endswith(".docx"):
        raw = _extract_text_from_docx(data)
        file_type = "Word"
    elif fname.endswith(".pdf"):
        raw = _extract_text_from_pdf(data)
        file_type = "PDF"
    elif fname.endswith(".txt") or fname.endswith(".csv"):
        raw = data.decode("utf-8", errors="replace")
        file_type = "텍스트"
    else:
        raise HTTPException(status_code=400,
            detail="지원 형식: .xlsx, .xls, .docx, .pdf, .txt, .csv")

    if len(raw.strip()) < 20:
        raise HTTPException(status_code=422,
            detail=f"{file_type} 파일에서 텍스트를 추출할 수 없습니다. 스캔 이미지 PDF는 지원되지 않습니다.")

    # 2. Excel 구조 파악 — 헤더 기반 직접 파싱 시도
    members = []
    if file_type == "Excel":
        members = _try_direct_excel_parse(data)

    # 3. 직접 파싱 실패 → Claude LLM 파싱
    if not members:
        if not ANTHROPIC_API_KEY:
            raise HTTPException(status_code=503,
                detail="ANTHROPIC_API_KEY가 없어 LLM 파싱을 수행할 수 없습니다.")
        client = _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        members = _parse_feedback_with_llm(client, raw, file.filename or "파일")

    return {
        "file_type": file_type,
        "filename":  file.filename,
        "member_count": len(members),
        "members": members,
        "source_preview": raw[:500],
    }


def _try_direct_excel_parse(data: bytes) -> list:
    """
    Excel 직접 파싱.
    ① 우선 표준 포맷 시도: 팀원 | 구분 | 항목 | 자기평가 | 팀장 피드백
       (구분: 본업/혁신/성장 — 동일 팀원 여러 행을 구분별로 합산)
    ② 단순 컬럼 포맷 시도: 이름/등급/본업/혁신/성장 헤더 기반
    """
    import io
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)

        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if len(rows) < 2:
                continue

            # ── 헤더 행 탐색 ──────────────────────────────────────
            header_idx = None
            for ri, row in enumerate(rows[:5]):
                cells = [str(c).strip() if c else "" for c in row]
                # 표준 포맷 감지: '팀원' + '구분' + '팀장 피드백' 컬럼 존재
                if any("팀원" in c for c in cells) and any("구분" in c for c in cells):
                    header_idx = ri
                    break
                # 단순 포맷 감지: '이름'/'성명' 컬럼 존재
                if any(c in ("이름","성명","팀원명","name") for c in cells):
                    header_idx = ri
                    break

            if header_idx is None:
                continue

            headers = [str(c).strip() if c else "" for c in rows[header_idx]]

            # ── 표준 포맷 (팀원|구분|항목|자기평가|팀장 피드백) ──
            def find_col(keywords):
                for ci, h in enumerate(headers):
                    if any(k in h for k in keywords):
                        return ci
                return -1

            col_member   = find_col(["팀원","이름","성명","name"])
            col_category = find_col(["구분","카테고리","category"])
            col_feedback = find_col(["팀장 피드백","팀장피드백","피드백","feedback"])

            if col_member >= 0 and col_category >= 0 and col_feedback >= 0:
                # 표준 포맷: 행별로 읽어서 팀원+구분으로 그룹핑
                members_dict: dict = {}
                order: list = []
                CAT_MAP = {
                    "본업": "job", "업무": "job", "job": "job",
                    "혁신": "innovation", "도전": "innovation", "창의": "innovation", "innovation": "innovation",
                    "성장": "growth", "역량": "growth", "개발": "growth", "학습": "growth", "growth": "growth",
                }
                for row in rows[header_idx + 1:]:
                    member = str(row[col_member] or "").strip()
                    if not member or member.lower() in ("none","nan",""):
                        continue
                    category_raw = str(row[col_category] or "").strip()
                    fb_text = str(row[col_feedback] or "").strip()

                    if member not in members_dict:
                        members_dict[member] = {"name": member, "grade": "", "job": [], "innovation": [], "growth": []}
                        order.append(member)

                    field = "job"  # 기본값
                    for key, mapped in CAT_MAP.items():
                        if key in category_raw:
                            field = mapped
                            break

                    if fb_text:
                        members_dict[member][field].append(fb_text)

                if members_dict:
                    return [
                        {
                            "name":       v["name"],
                            "grade":      v["grade"],
                            "job":        "\n\n".join(v["job"]),
                            "innovation": "\n\n".join(v["innovation"]),
                            "growth":     "\n\n".join(v["growth"]),
                        }
                        for v in (members_dict[n] for n in order)
                    ]

            # ── 단순 컬럼 포맷 (이름/등급/본업/혁신/성장 한 행 = 한 명) ──
            KEY_MAP = {
                "name":       ["이름","성명","팀원명","팀원","name"],
                "grade":      ["등급","평가등급","grade"],
                "job":        ["본업","담당업무","업무","job"],
                "innovation": ["혁신","도전","혁신","innovation"],
                "growth":     ["성장","역량","성장","growth"],
            }
            col_map: dict = {}
            for field, keys in KEY_MAP.items():
                for ci, h in enumerate(headers):
                    if any(k in h for k in keys):
                        col_map[field] = ci
                        break

            if "name" not in col_map:
                continue

            members = []
            for row in rows[header_idx + 1:]:
                name = str(row[col_map["name"]] or "").strip()
                if not name or name.lower() in ("none","nan",""):
                    continue
                members.append({
                    "name":       name,
                    "grade":      str(row[col_map["grade"]] or "").strip() if "grade" in col_map else "",
                    "job":        str(row[col_map["job"]] or "").strip() if "job" in col_map else "",
                    "innovation": str(row[col_map["innovation"]] or "").strip() if "innovation" in col_map else "",
                    "growth":     str(row[col_map["growth"]] or "").strip() if "growth" in col_map else "",
                })
            if members:
                return members

    except Exception:
        pass
    return []


# ── Ellis: Word 리포트 생성 ──────────────────────────────────

class EllisReportBody(BaseModel):
    manager_name: str = "팀장"
    analysis:     dict


@app.post("/api/ellis/generate-docx")
async def ellis_generate_docx(body: EllisReportBody):
    """분석 결과를 Word(.docx)로 변환하여 반환"""
    import io
    from fastapi.responses import StreamingResponse
    try:
        from docx import Document as DocxDoc
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=503, detail="python-docx 미설치")

    a     = body.analysis or {}
    sum_  = a.get("summary", {})
    dims  = a.get("dimensions", {})
    mems  = a.get("member_analysis", [])
    acts  = a.get("action_items", [])
    today_s = datetime.now().strftime("%Y년 %m월 %d일")

    # ── 안전한 헬퍼 ──────────────────────────────────────────
    def add_colored(doc, text: str, color: RGBColor, size_pt: int = 10):
        """텍스트가 있을 때만 단락 추가 + 색상 적용"""
        if not text:
            return
        p = doc.add_paragraph()
        run = p.add_run(str(text))
        run.font.color.rgb = color
        run.font.size = Pt(size_pt)

    def add_plain(doc, text: str, size_pt: int = 10):
        if not text:
            return
        p = doc.add_paragraph()
        run = p.add_run(str(text))
        run.font.size = Pt(size_pt)

    # ── 문서 작성 ─────────────────────────────────────────────
    doc = DocxDoc()

    # 제목
    title_p = doc.add_heading("Ellis in Wonderland — Feedback Analysis", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f"작성일: {today_s}   |   By Agent Ellis")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x7A, 0x75, 0x7F)
    doc.add_paragraph()

    # 1. 종합 요약
    doc.add_heading("1. 종합 요약", 1)
    score = sum_.get("overall_score", "—")
    grade = sum_.get("grade", "—")
    add_plain(doc, f"전체 점수: {score} / 100   등급: {grade}", 11)
    add_plain(doc, f"헤드라인: {sum_.get('headline','—')}", 10)
    add_plain(doc, f"잘된 점: {sum_.get('top_strength','—')}", 10)
    add_plain(doc, f"주의 점: {sum_.get('top_risk','—')}", 10)
    doc.add_paragraph()

    # 2. 5개 축 분석
    doc.add_heading("2. 5개 축 분석", 1)
    DIM_KO = {
        "objectivity": "객관성",
        "fairness":    "공정성",
        "balance":     "균형성",
        "consistency": "일관성",
        "culture_fit": "문화 적합성",
    }
    for k, ko in DIM_KO.items():
        d = dims.get(k, {})
        doc.add_heading(f"{ko}: {d.get('score','—')}점", 2)
        add_plain(doc, d.get("finding", ""), 10)

        risks = d.get("risk_expressions") or d.get("high_risk_expressions") or []
        if risks:
            add_colored(doc, "주의 표현: " + " / ".join(str(r) for r in risks),
                        RGBColor(0x9b, 0x1b, 0x1b))

        goods = d.get("good_examples") or d.get("coaching_suggestions") or []
        if goods:
            add_colored(doc, "긍정·권고: " + " / ".join(str(g) for g in goods),
                        RGBColor(0x0d, 0x6b, 0x47))

        if d.get("outlier_members"):
            add_colored(doc, "재검토 팀원: " + ", ".join(str(x) for x in d["outlier_members"]),
                        RGBColor(0xa1, 0x64, 0x00))
        doc.add_paragraph()

    # 3. 팀원별 분석
    doc.add_heading("3. 팀원별 분석 (익명화됨)", 1)
    for m in mems:
        name = m.get("name", "팀원")
        doc.add_heading(name, 2)

        sc = m.get("scores", {})
        if sc:
            add_plain(doc, "분석 점수: " + "  /  ".join(f"{k} {v}" for k, v in sc.items()))

        flags = m.get("flags", [])
        if flags:
            add_colored(doc, "리스크: " + ", ".join(str(f) for f in flags),
                        RGBColor(0x9b, 0x1b, 0x1b))

        if m.get("strength"):
            add_colored(doc, f"강점: {m['strength']}", RGBColor(0x0d, 0x6b, 0x47))
        if m.get("improvement"):
            add_colored(doc, f"개선 제안: {m['improvement']}", RGBColor(0xa1, 0x64, 0x00))
        doc.add_paragraph()

    # 4. 개선 권고사항
    if acts:
        doc.add_heading("4. 개선 권고사항", 1)
        for i, ac in enumerate(acts, 1):
            priority = ac.get("priority", "—")
            item     = ac.get("item", "")
            color = RGBColor(0x9b, 0x1b, 0x1b) if priority == "높음" else RGBColor(0xa1, 0x64, 0x00)
            add_colored(doc, f"{i}. [{priority}] {item}", color, 10)

    # 면책 주석
    doc.add_paragraph()
    add_colored(doc,
        "※ 본 리포트는 피드백 패턴 관찰을 위한 참고 자료입니다. 팀원 식별정보는 포함되지 않습니다.",
        RGBColor(0x9E, 0x9E, 0x9E), 8)

    # ── 버퍼에 저장 후 스트리밍 반환 ─────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"Ellis_FeedbackAnalysis_{datetime.now().strftime('%Y%m%d')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'}
    )


# ── 엔드포인트: 납입고지서 OCR ───────────────────────────────

class OcrBody(BaseModel):
    image_base64: str = Field(..., description="base64 인코딩된 이미지 (PNG/JPG)")
    filename: str = Field("", description="원본 파일명 (미디어 타입 추론용)")


@app.post("/api/ocr-invoice")
async def ocr_invoice(body: OcrBody):
    """
    납입고지서 이미지를 Claude Vision 으로 분석하여 구조화된 JSON 반환.
    프론트에서 PDF → canvas → base64 PNG 로 변환 후 전송.
    """
    client = get_client()

    # 미디어 타입 추론
    fname = (body.filename or "").lower()
    if fname.endswith(".png") or body.image_base64.startswith("iVBOR"):
        media_type = "image/png"
    elif fname.endswith(".gif"):
        media_type = "image/gif"
    elif fname.endswith(".webp"):
        media_type = "image/webp"
    else:
        media_type = "image/jpeg"

    prompt = """이 이미지는 한국의 교육비 납입 서류(교육비납입확인서, 학자금납입확인서, 학자금납입증명서 등)입니다.
한국어로 작성된 서류에서 아래 정보를 추출하여 JSON만 반환하세요. 다른 텍스트는 출력하지 마세요.
확인 불가 필드는 null로 표시하세요.

[서류 유형별 필드 위치]
1. 교육비납입확인서(유치원/어린이집): 원아명→child_name, 보호자명→guardian_name, 합계금액→amount_total
2. 학자금납입확인서(고등학교): 학생명→child_name, 보호자명→guardian_name, 합계→amount_total
3. 학자금납입증명서(대학교): 성명→child_name, 학번→student_id, 합계→amount_total
4. 국제학교/해외: school_country에 ISO 국가코드(US/UK/SG 등)
5. "영수증" 제목의 학교 발급 서류: doc_type="학자금납입확인서"로 매핑. 발행기관이 "○○학교장"이면 학교급 추론.

[영수증 형식 - 매우 중요한 가로 4칸 병렬 표 해석]
"영 수 증" 제목의 학교 발급 서류는 아래처럼 같은 행에 학번·성명이 병렬 배치됩니다:
  [학 번 | 학년·반·번호 정보 | 성 명 | 학생이름]
  [금 액 |                금액 원 정              ]
  [내 역 | 차 수 | 금 액 | 이체일자]
  [수업료 | 1/4분기 | 1,360,000 | 2026-03-26]
-> "성 명" 셀의 오른쪽 옆 셀 값이 child_name (반드시 추출)
-> "학 번" 셀의 오른쪽 옆 셀 값에 학년·반·번호 -> grade 필드
-> "수업료" 행의 "이체일자" 열 값 -> payment_date
-> "계/합계" 행의 "금액" 열 값 -> amount_total
-> 발행기관 "인천하늘고등학교장" -> school_name="인천하늘고등학교", school_type="고등학교"

[child_name 추출 규칙 - 매우 중요, 절대 간과 금지]
학생/피부양자 이름은 아래 라벨 중 어떤 것이든 등장하면 반드시 child_name으로 매핑하세요.
허용 라벨(공백·줄바꿈 무관, 대소문자 무관):
  - "성명", "성 명", "성  명", "성\t명"  (가장 흔함 - 표 형식에서 글자 사이 공백이 있을 수 있음)
  - "학생명", "학생 성명", "학생성명", "학생 명"
  - "원아명", "유아명", "자녀명", "자녀 성명"
  - "피부양자 성명", "피보험자 성명", "수혜자명", "수혜자 성명"
  - "납입자명"(피교육자 본인이 대학생 납부자인 경우)
  - 영문: "Name", "Student Name", "Full Name"

child_name vs guardian_name 구분 원칙 (우선순위 순):
  1) 문서에 "학생정보/수혜자정보/피부양자정보" 섹션이 있으면 그 안의 성명 = child_name
  2) "납부자정보/보호자정보/학부모" 섹션의 성명 = guardian_name (child_name 아님)
  3) 성명 옆이나 같은 행에 학번/학년/반/학과/전공/생년월일이 있으면 = child_name
  4) 성명 옆에 사번/직원번호/부서/소속회사가 있으면 = guardian_name (회사 직원)
  5) 학자금납입증명서처럼 피교육자 본인 명의 서류면 문서 상단의 성명 = child_name
  6) 성명이 하나만 있고 구분이 애매하면 child_name으로 우선 매핑

절대 null 금지 조건:
  이미지에 한글 2~4자의 사람 이름으로 보이는 문자열이 "성명" 계열 라벨과 함께 나타나면
  반드시 그 값을 child_name으로 반환하세요. 라벨과 이름 사이 공백/줄바꿈은 무시하세요.
  표 형식에서 "성", "명"이 세로나 가로로 분리돼 있어도 하나의 라벨로 해석하세요.

반환 JSON:
{
  "doc_type": "교육비납입확인서|학자금납입확인서|학자금납입증명서",
  "child_name": "학생/원아 이름",
  "guardian_name": "보호자/학부모 이름 또는 null",
  "student_id": "학번(대학교만) 또는 null",
  "major": "학과(대학교만) 또는 null",
  "grade": "학년반(초중고/유치원) 또는 null",
  "school_name": "학교/유치원/기관 이름만 (원장·교장·총장 등 직함·인명 제외. 예: '미래유치원장' → '미래유치원')",
  "school_type": "유치원|초등학교|중학교|고등학교|대학교",
  "school_country": "KR(기본값, 해외면 US/UK/SG 등)",
  "school_address": "주소(있는 경우)",
  "issuer_name": "교장/원장/총장 이름과 직함 (school_name과 분리하여 여기에만 기재)",
  "issue_date": "발급일 YYYY-MM-DD",
  "payment_date": "납부일 YYYY-MM-DD(여러 건이면 가장 최근)",
  "benefit_year": 연도정수,
  "benefit_term": "1분기~4분기 또는 1학기/2학기",
  "amount_tuition": 수업료/등록금_정수,
  "amount_ops": 운영지원비/특별활동비_정수(없으면 0),
  "amount_other": 기타항목합계_정수(없으면 0),
  "amount_total": 납부총액_정수,
  "confidence": 0.0~1.0,
  "notes": "특이사항(직인확인여부, 복수월납부 등)"
}"""

    try:
        message = client.messages.create(
            model=OCR_MODEL,         # OCR은 Sonnet 4.6 - Haiku 대비 표·세로라벨 해석 정확도 우선
            max_tokens=2048,         # Sonnet 대응 충분 여유 (기존 1024는 긴 서류에서 truncation 위험)
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": body.image_base64,
                        },
                    },
                    {"type": "text", "text": prompt},
                ],
            }],
        )
        raw = message.content[0].text if message.content else "{}"
    except _anthropic.AuthenticationError:
        raise HTTPException(status_code=401, detail="ANTHROPIC_API_KEY 가 유효하지 않습니다.")
    except _anthropic.RateLimitError:
        raise HTTPException(status_code=429, detail="Anthropic API 한도 초과. 잠시 후 재시도하세요.")
    except _anthropic.APIStatusError as e:
        raise HTTPException(status_code=e.status_code, detail=f"Anthropic API 오류: {e.message}")

    json_match = re.search(r"\{[\s\S]*\}", raw)
    raw_json = json_match.group(0) if json_match else raw
    try:
        result = json.loads(raw_json)
    except json.JSONDecodeError:
        raise HTTPException(status_code=422, detail=f"OCR 결과 파싱 실패: {raw[:300]}")

    # ── child_name 후처리 fallback ──
    # 프롬프트가 null을 돌려준 경우에도 raw 응답 텍스트에서 성명 패턴을 직접 재추출.
    # 한글 2~4자 이름만 수용. 흔한 거짓 양성(예: "성명란", "성명표")은 자체 필터.
    def _recover_child_name(text: str) -> str | None:
        if not text:
            return None
        # notes/raw에 "홍길동" 류 이름이 라벨과 함께 등장하는 모든 변형을 포착
        patterns = [
            r'학\s*생\s*(?:성\s*)?명\s*[:：\s]*([가-힣]{2,4})(?![가-힣])',
            r'원\s*아\s*명\s*[:：\s]*([가-힣]{2,4})(?![가-힣])',
            r'자\s*녀\s*명\s*[:：\s]*([가-힣]{2,4})(?![가-힣])',
            r'피\s*부\s*양\s*자\s*성\s*명\s*[:：\s]*([가-힣]{2,4})(?![가-힣])',
            r'수\s*혜\s*자\s*(?:성\s*)?명\s*[:：\s]*([가-힣]{2,4})(?![가-힣])',
            r'성\s*명\s*[:：\s]+([가-힣]{2,4})(?![가-힣])',
        ]
        _blacklist = {"성명란", "성명표", "성명서", "성명록", "성명확인"}
        for pat in patterns:
            for m in re.finditer(pat, text):
                candidate = m.group(1).strip()
                if candidate and candidate not in _blacklist:
                    return candidate
        return None

    if not (result.get("child_name") or "").strip():
        # 1차: notes 필드에 모델이 남긴 언급 확인
        recovered = _recover_child_name(result.get("notes") or "")
        if not recovered:
            # 2차: raw 전체에서 재시도 (모델이 주석처럼 남긴 텍스트도 포함)
            recovered = _recover_child_name(raw)
        if recovered:
            result["child_name"] = recovered
            result["_child_name_source"] = "regex_fallback"

    return result


# ══════════════════════════════════════════════════════════════
# HR 뉴스룸 API  (/api/hr-newsroom/digest)
# ══════════════════════════════════════════════════════════════

class NewsroomBody(BaseModel):
    morning_days:     int = Field(2,  ge=1, le=7)
    education_days:   int = Field(30, ge=7, le=90)
    conference_days:  int = Field(21, ge=7, le=60)


def _llm_generate_edu_conf(client: "_anthropic.Anthropic") -> dict:
    """
    Claude가 HR 교육 추천(가격·일정·기관 포함)과
    글로벌·국내 컨퍼런스 상세 정보를 직접 생성.
    """
    from datetime import timedelta
    today = datetime.now()
    today_str   = today.strftime("%Y년 %m월 %d일")
    date_from   = (today + timedelta(weeks=3)).strftime("%Y년 %m월 %d일")
    date_to     = (today + timedelta(days=60)).strftime("%Y년 %m월 %d일")

    prompt = f"""당신은 한화투자증권 혁신지원실에서 인사팀 직원들의 역량개발을 지원하는 전문가입니다.
현재 날짜: {today_str}

⚠️ 중요 제약: 교육과 컨퍼런스의 일정은 반드시 **{date_from} ~ {date_to}** 사이에 시작하는 것만 포함하세요.
(현재 날짜 기준 3주 후 ~ 2개월 후 범위. 이 기간 밖의 일정은 절대 포함하지 마세요.)

아래 두 가지를 JSON 형식으로 작성해주세요. JSON 외 다른 텍스트 없이 출력하세요.

1. **인사팀 추천 역량개발 교육** (6건)
반드시 아래 기관들을 우선 포함하여 다양하게 구성하세요:
- 중앙경제교육원 (중앙일보 계열, 비즈니스·HR 교육)
- 한국생산성본부 (KPC, 경영·HRD·디지털 역량)
- 한국HR협의회 / 한국HRD협회 (인사·채용·평가 전문)
- 노사발전재단 (노무·노동법 실무)
- 휴넷 또는 패스트캠퍼스 (온라인 직무교육)
- SHRM Korea 또는 한국경영자총협회 (글로벌 HR 트렌드)

각 교육 필드:
- title: 교육 과정명 (구체적이고 실무적인 제목)
- organizer: 주최 기관명
- format: "온라인" 또는 "오프라인"
- price: 가격 (예: "350,000원", "무료", "550,000원 (조기등록 450,000원)")
- schedule: 일정 (예: "2026년 5월 8일~9일 (2일)", "매월 1·3주 화요일", "상시 수강")
- duration: 교육 시간 또는 기간
- location: 장소 (오프라인: 건물명+지역 / 온라인: "Zoom 실시간" 또는 "LMS 자기주도")
- recommended_for: 추천 대상 (예: "채용 담당자", "노무 담당자", "인사기획 팀장", "전 인사팀")
- reason: 추천 이유 2~3줄 — 왜 지금 이 교육이 인사팀에 필요한지 구체적으로
- url: 해당 기관의 실제처럼 보이는 URL

2. **HR·디지털·핀테크 컨퍼런스** (국내 4건 + 글로벌 4건, 총 8건)
주제는 HR, AI·디지털 혁신, 핀테크·금융 테크를 다양하게 포함하세요.
2026년 상반기~하반기 기준의 실제 개최될 법한 행사.

각 컨퍼런스 필드:
- title: 컨퍼런스명
- type: "국내" 또는 "글로벌"
- category: "HR" 또는 "디지털·AI" 또는 "핀테크·금융"
- organizer: 주최 기관
- date: 일정 (예: "2026년 6월 11일~12일")
- venue: 개최 장소 (건물명·홀명, 도시, 국가)
- price: 참가비 (예: "무료", "일반 180,000원 / 단체(5인↑) 150,000원", "USD 1,500 / 얼리버드 USD 1,100")
- language: "한국어" 또는 "영어" 또는 "영어 (한국어 동시통역 제공)"
- topics: 주요 세션 주제 배열 (3~4개)
- description: 행사 소개 1~2문장
- recommend_reason: 인사팀이 참가해야 하는 이유 1~2문장
- url: 실제처럼 보이는 URL

국내 컨퍼런스 예시 방향 (이 외에도 추가):
- HR 관련: 한국HR컨퍼런스, 인사혁신포럼
- 디지털·AI: 디지털이노베이션코리아, AI EXPO Korea
- 핀테크·금융: 코리아핀테크위크, 금융IT이노베이션포럼

글로벌 컨퍼런스 예시 방향:
- HR: SHRM Annual, HR Tech Conference (Las Vegas)
- 디지털·AI: Gartner IT Symposium, Google Cloud Next
- 핀테크: Money20/20, Sibos

JSON 형식:
{{
  "education": [...6건...],
  "conferences": [...8건...]
}}"""

    try:
        msg = client.messages.create(
            model=CHAT_MODEL, max_tokens=4000,
            messages=[{"role": "user", "content": prompt}]
        )
        raw = msg.content[0].text if msg.content else "{}"
        m = re.search(r"\{[\s\S]*\}", raw)
        return json.loads(m.group(0)) if m else {}
    except Exception:
        return {}


@app.post("/api/hr-newsroom/digest")
async def hr_newsroom_digest(body: NewsroomBody):
    """
    Google News RSS 수집 → Claude LLM 큐레이션 → 구조화된 JSON 반환.
    LLM이 무관한 기사를 제거하고 HR 실무 관련성을 평가·요약합니다.
    """
    from collections import defaultdict

    # ── 1. RSS 수집 (병렬) ───────────────────────────────────
    import asyncio
    labor_raw, recruit_raw, edu_raw, edu_rec_raw, conf_raw = await asyncio.gather(
        _fetch_google_news_rss("고용노동부 OR 포괄임금 OR 노동법 OR 근로기준법 OR 임금체불 OR 노사관계", days=body.morning_days, limit=10),
        _fetch_google_news_rss("경력채용 OR 신입채용 OR 채용트렌드 OR 인재확보 OR 채용시장", days=body.morning_days, limit=10),
        _fetch_google_news_rss("기업교육 OR HRD OR 역량개발 OR 인사관리 OR HR트렌드 OR 조직문화", days=body.morning_days, limit=10),
        _fetch_google_news_rss("직무교육 OR 리더십교육 OR 온라인교육 OR HR교육 OR 역량강화", days=body.education_days, limit=15),
        _fetch_google_news_rss("AI 컨퍼런스 OR AI conference OR 디지털혁신 포럼 OR LLM 세미나 OR GenAI", days=body.conference_days, limit=20),
    )

    import asyncio

    # ── 2a. 조간 브리핑용 RSS 수집 (병렬) ────────────────────
    labor_raw, recruit_raw, edu_raw = await asyncio.gather(
        _fetch_google_news_rss("고용노동부 OR 포괄임금 OR 노동법 OR 근로기준법 OR 임금체불", days=body.morning_days, limit=10),
        _fetch_google_news_rss("경력채용 OR 신입채용 OR 채용트렌드 OR 인재확보", days=body.morning_days, limit=10),
        _fetch_google_news_rss("기업교육 OR HRD OR 역량개발 OR 인사관리 OR HR트렌드", days=body.morning_days, limit=10),
    )

    # ── 2b. Claude로 브리핑 큐레이션 ─────────────────────────
    morning_sections: dict = {"노무": [], "채용": [], "교육": []}
    llm_curated = False

    if ANTHROPIC_API_KEY:
        try:
            client_llm = _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

            def items_to_text(items: list) -> str:
                return "\n".join(
                    f"{i+1}. {it['title']} ({it.get('source','')}, {it.get('published_at','')[:10]})"
                    for i, it in enumerate(items)
                )

            brief_prompt = f"""HR 전문가로서 아래 뉴스 후보에서 인사·노무·채용 실무에 실제로 유용한 기사를 선별하세요.
홍보성·단순인사이동·무관 경제기사는 제외합니다.
각 카테고리에서 3건 선택, 한 줄 요약(30자이내)과 관련성점수(1-10) 포함.
JSON만 출력 (마크다운 없이):
{{"노무":[{{"index":번호,"summary":"요약","relevance":점수}}],"채용":[...],"교육":[...]}}

[노무후보]\n{items_to_text(labor_raw)}\n[채용후보]\n{items_to_text(recruit_raw)}\n[교육후보]\n{items_to_text(edu_raw)}"""

            msg = client_llm.messages.create(model=CHAT_MODEL, max_tokens=800,
                messages=[{"role":"user","content":brief_prompt}])
            raw_j = msg.content[0].text if msg.content else "{}"
            m = re.search(r"\{[\s\S]*\}", raw_j)
            curated = json.loads(m.group(0)) if m else {}

            def apply(raw_list, picks, n=3):
                res = []
                for p in picks[:n]:
                    idx = int(p.get("index",0))-1
                    if 0 <= idx < len(raw_list):
                        it = dict(raw_list[idx])
                        it["summary"] = p.get("summary","")
                        it["relevance"] = p.get("relevance", 5)
                        res.append(it)
                return res or raw_list[:n]

            morning_sections = {
                "노무": apply(labor_raw,   curated.get("노무",[])),
                "채용": apply(recruit_raw, curated.get("채용",[])),
                "교육": apply(edu_raw,     curated.get("교육",[])),
            }
            llm_curated = True

            # ── 2c. 교육·컨퍼런스 상세 정보 생성 ────────────
            edu_conf = _llm_generate_edu_conf(client_llm)

        except Exception:
            for cat, raw in [("노무",labor_raw),("채용",recruit_raw),("교육",edu_raw)]:
                morning_sections[cat] = raw[:3]
            edu_conf = {}
    else:
        for cat, raw in [("노무",labor_raw),("채용",recruit_raw),("교육",edu_raw)]:
            morning_sections[cat] = raw[:3]
        edu_conf = {}

    # ── 3. 교육·컨퍼런스 폴백 (LLM 실패 시) ─────────────────
    edu_items = edu_conf.get("education", []) or _get_fallback_education()
    conf_items_raw = edu_conf.get("conferences", []) or _get_fallback_conferences()

    # 컨퍼런스: 국내/글로벌 분리
    conf_domestic = [c for c in conf_items_raw if c.get("type") == "국내"]
    conf_global   = [c for c in conf_items_raw if c.get("type") == "글로벌"]

    return {
        "generated_at": datetime.now().isoformat(),
        "llm_curated": llm_curated,
        "morning_brief": {"sections": morning_sections},
        "education_recommendations": {"items": edu_items},
        "conferences": {
            "domestic": conf_domestic,
            "global":   conf_global,
            "all":      conf_items_raw,
        },
        # 하위호환
        "ai_conference_daily": {
            "timeline": [
                {"date": c.get("date",""), "items": [{"title":c.get("title",""), "url":c.get("url","#"), "source":c.get("organizer","")}]}
                for c in conf_items_raw
            ]
        },
    }


def _get_fallback_education() -> list:
    """LLM 실패 시 기본 교육 목록 (3주~2개월 후 기준: 5월 초~6월 중순)"""
    return [
        {"title":"2026 인사·노무 실무 종합과정","organizer":"중앙경제교육원","format":"오프라인","price":"520,000원 (조기등록 430,000원)","schedule":"2026년 5월 12일~14일 (3일)","duration":"21시간","location":"서울 중구 중앙빌딩 교육센터","recommended_for":"인사·노무 담당자","reason":"중앙경제교육원 대표 HR 과정. 노동법 개정사항·임금관리·인사평가 실무를 집중 학습합니다. 금융권 HR 실무자 중심 사례 포함.","url":"https://www.junganghr.com/edu/2026-hr-master"},
        {"title":"노동법 최신 판례 및 징계·해고 실무","organizer":"노사발전재단","format":"온라인","price":"무료 (고용노동부 지원)","schedule":"2026년 5월 7일~28일 (4주, Zoom)","duration":"20시간","location":"Zoom 실시간 강의","recommended_for":"노무 담당자·인사팀장","reason":"2025~2026년 대법원 판례·고용노동부 지침 변경사항을 정리합니다. 무료 정부지원 과정으로 신청 경쟁률이 높아 조기 마감될 수 있습니다.","url":"https://www.nosa.or.kr/edu/labor-law-2026"},
        {"title":"HRD 전략기획 전문가 과정","organizer":"한국생산성본부 (KPC)","format":"오프라인","price":"450,000원","schedule":"2026년 5월 20일~21일 (2일)","duration":"16시간","location":"서울 여의도 KPC교육센터","recommended_for":"HRD 기획·교육팀","reason":"KPC 핵심 HRD 과정. 교육체계 설계·역량모델링·교육 ROI 측정까지 다룹니다. 금융권 HRD 담당자에게 특히 유용합니다.","url":"https://www.kpc.or.kr/edu/hrd-strategic-2026"},
        {"title":"AI 기반 채용·평가 혁신 실무","organizer":"한국HR협의회","format":"오프라인","price":"380,000원","schedule":"2026년 6월 4일 (1일)","duration":"8시간","location":"서울 강남 HR협의회 세미나실","recommended_for":"채용 담당자·HR BP","reason":"AI 역량검사 도구 비교·구조화 면접 설계·데이터 기반 채용 의사결정 방법론을 실습합니다. 금융권 채용 특화 케이스 포함.","url":"https://www.hrkorea.or.kr/edu/ai-recruit-2026"},
        {"title":"성과관리·OKR 도입 실전 워크숍","organizer":"한국HRD협회","format":"오프라인","price":"350,000원","schedule":"2026년 6월 11일~12일 (2일)","duration":"14시간","location":"서울 마포 HRD협회 교육관","recommended_for":"인사기획·성과관리 담당자","reason":"OKR 설계부터 전사 성과관리 체계·평가 보정 방법론까지 다룹니다. 금융·증권사 적용 실제 사례 포함.","url":"https://www.hrd.or.kr/workshop/okr-2026"},
        {"title":"디지털 HR 전환 리더십 과정","organizer":"한국생산성본부 (KPC)","format":"온라인","price":"280,000원","schedule":"2026년 5월 11일 개강 (6주 자기주도)","duration":"30시간","location":"KPC 온라인 LMS","recommended_for":"인사팀장·HR 리더","reason":"HR 디지털 전환 전략·피플애널리틱스 기초·AI HR 도구 활용법을 체계적으로 학습합니다. 바쁜 관리자를 위한 자기주도 과정.","url":"https://www.kpc.or.kr/elearning/digital-hr-2026"},
    ]


def _get_fallback_conferences() -> list:
    """LLM 실패 시 기본 컨퍼런스 목록 — 3주~2개월 후 범위 (5월 초~6월 중순) 내 일정 우선"""
    return [
        # 국내 4건
        {"title":"2026 대한민국 인사혁신 컨퍼런스","type":"국내","category":"HR","organizer":"한국경영자총협회","date":"2026년 6월 11일~12일","venue":"코엑스 그랜드볼룸, 서울","price":"일반 180,000원 / 단체(5인↑) 140,000원","language":"한국어","topics":["AI HR","성과관리","조직문화","채용혁신"],"description":"국내 최대 HR 컨퍼런스. 주요 기업 인사 담당자 2,000명 이상 참가.","recommend_reason":"국내 HR 최신 트렌드와 우수 기업 사례를 한자리에서 파악할 수 있습니다.","url":"https://www.kef.or.kr/conf/2026-hr"},
        {"title":"2026 HR·피플애널리틱스 포럼","type":"국내","category":"HR","organizer":"한국HR협의회","date":"2026년 5월 22일","venue":"서울 여의도 전경련 컨퍼런스센터","price":"일반 120,000원 / 회원 80,000원","language":"한국어","topics":["피플애널리틱스","HR 데이터 활용","AI 인재관리","채용 데이터"],"description":"데이터 기반 HR 의사결정 전문 포럼. 국내 주요 기업 HR 데이터 활용 사례 발표.","recommend_reason":"피플애널리틱스 도입 사례와 데이터 기반 채용·평가 방법론을 직접 배울 수 있습니다.","url":"https://www.hrkorea.or.kr/forum/analytics-2026"},
        {"title":"금융IT이노베이션포럼 2026","type":"국내","category":"핀테크·금융","organizer":"금융결제원·금융보안원","date":"2026년 5월 29일","venue":"서울 중구 은행회관 국제회의실","price":"무료 (사전등록 필수)","language":"한국어","topics":["AI 금융서비스","클라우드 보안","디지털 뱅킹","RegTech"],"description":"국내 금융권 IT·보안·혁신 담당자 대상 연례 포럼.","recommend_reason":"금융 디지털 전환 트렌드를 파악하고 금융IT 인재 채용 방향 수립에 활용할 수 있습니다.","url":"https://www.kftc.or.kr/forum/finit-2026"},
        {"title":"2026 디지털·AI 인재 채용 전략 세미나","type":"국내","category":"디지털·AI","organizer":"중앙경제교육원","date":"2026년 6월 5일","venue":"서울 중구 중앙빌딩 대강당","price":"일반 95,000원 / 조기등록 70,000원","language":"한국어","topics":["디지털 인재 발굴","AI 직무역량 검증","기술 인재 채용 전략","MZ세대 조직문화"],"description":"금융·IT 업계 디지털·AI 인재 채용 트렌드와 실전 전략 공유.","recommend_reason":"디지털 전환 가속화에 따른 기술인재 채용 전략을 최신 사례 중심으로 학습할 수 있습니다.","url":"https://www.junganghr.com/seminar/digital-recruit-2026"},
        # 글로벌 4건
        {"title":"SHRM Annual Conference & Expo 2026","type":"글로벌","category":"HR","organizer":"SHRM","date":"2026년 6월 21일~25일","venue":"McCormick Place, 시카고, 미국","price":"USD 2,295 (얼리버드 USD 1,895)","language":"영어","topics":["Future of Work","People Analytics","DEI","AI in HR"],"description":"세계 최대 HR 전문가 컨퍼런스. 전 세계 140개국 2만명 이상 참가.","recommend_reason":"글로벌 HR 트렌드와 최신 HR 테크 솔루션을 직접 경험할 수 있습니다.","url":"https://annual.shrm.org/2026"},
        {"title":"ATD International Conference & Exposition 2026","type":"글로벌","category":"HR","organizer":"ATD (Association for Talent Development)","date":"2026년 5월 17일~20일","venue":"Washington Convention Center, 워싱턴 D.C., 미국","price":"USD 1,995 (얼리버드 USD 1,595)","language":"영어","topics":["Learning & Development","Talent Management","Instructional Design","AI in L&D"],"description":"전 세계 HRD·교육 전문가 10,000명 이상 참가하는 글로벌 최대 HRD 컨퍼런스.","recommend_reason":"HRD 최신 방법론과 AI 기반 교육 솔루션을 직접 체험하고 글로벌 HRD 전문가와 네트워킹할 수 있습니다.","url":"https://atdconference.td.org/2026"},
        {"title":"Money20/20 Europe 2026","type":"글로벌","category":"핀테크·금융","organizer":"Money20/20","date":"2026년 6월 2일~4일","venue":"RAI Amsterdam Convention Centre, 암스테르담, 네덜란드","price":"EUR 2,895 (얼리버드 EUR 2,295)","language":"영어","topics":["Open Banking","AI in Finance","Digital Payments","Embedded Finance"],"description":"유럽 최대 핀테크 컨퍼런스. 6,000명 이상 참가, 300개국 핀테크 리더 집결.","recommend_reason":"글로벌 금융 혁신 트렌드를 파악하고 핀테크 분야 인재 채용 동향을 선제적으로 확인할 수 있습니다.","url":"https://europe.money2020.com/2026"},
        {"title":"Google Cloud Next '26","type":"글로벌","category":"디지털·AI","organizer":"Google","date":"2026년 5월 19일~21일","venue":"Moscone Center, 샌프란시스코, 미국","price":"USD 1,299 (일반 / 얼리버드 USD 899)","language":"영어 (주요 세션 한국어 자막)","topics":["Gemini AI","Cloud AI","Workspace AI","Enterprise AI 도입"],"description":"Google Cloud 연례 최대 행사. 전 세계 3만명 이상 참가. 최신 AI·클라우드 발표.","recommend_reason":"AI 도구 기업 도입 사례를 통해 HR 업무 자동화 방향성과 디지털 인재 수요를 파악할 수 있습니다.","url":"https://cloud.withgoogle.com/next/2026"},
    ]


# ── 엔드포인트: URL 본문 가져오기 ────────────────────────────

class FetchBody(BaseModel):
    url: str = Field(..., min_length=8)


@app.post("/api/fetch-reference")
async def fetch_reference(body: FetchBody):
    url = body.url.strip()
    if not url.startswith(("http://", "https://")):
        raise HTTPException(status_code=400, detail="http(s) URL만 지원합니다.")
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=20.0) as hx:
            r = await hx.get(
                url,
                headers={
                    "User-Agent": "Mozilla/5.0 (compatible; HR-JD-Bot/1.0; +local)",
                    "Accept": "text/html,application/xhtml+xml;q=0.9,*/*;q=0.8",
                },
            )
            r.raise_for_status()
            ctype = r.headers.get("content-type", "")
            raw = r.text
    except httpx.HTTPError as e:
        raise HTTPException(status_code=502, detail=f"URL 가져오기 실패: {e!s}") from e

    title = ""
    if "html" in ctype.lower():
        title_match = re.search(r"<title[^>]*>([^<]+)</title>", raw, re.I)
        if title_match:
            title = BeautifulSoup(title_match.group(1), "html.parser").get_text(strip=True)
        text = strip_html(raw)
    else:
        text = raw

    text = text.strip()[:200_000]
    if len(text) < 80:
        raise HTTPException(
            status_code=422,
            detail="추출된 본문이 너무 짧습니다. 다른 URL을 시도하거나 JD를 직접 붙여 넣어 주세요.",
        )
    return {"title": title or url, "text": text, "url": url}


# ── 엔드포인트: 이력서 일괄 분석 ─────────────────────────────

class ResumeIn(BaseModel):
    filename: str
    text: str


class ScoreBatchBody(BaseModel):
    jd_text: str = Field("", description="직무기술서(JD) 또는 공고 전문")
    resumes: list[ResumeIn]
    top_k: int = Field(default=TOP_K, ge=1, le=12)


@app.post("/api/score-resumes")
async def score_resumes(body: ScoreBatchBody):
    jd = (body.jd_text or "").strip()
    if len(jd) < 120:
        raise HTTPException(
            status_code=400,
            detail="직무 레퍼런스(JD/공고 텍스트)가 너무 짧습니다. URL 불러오기 후 본문을 보강하거나 직접 붙여 주세요.",
        )
    if not body.resumes:
        raise HTTPException(status_code=400, detail="이력서가 없습니다.")

    client = get_client()
    chunks = chunk_text(jd)
    if not chunks:
        raise HTTPException(status_code=400, detail="JD에서 유효한 청크를 만들 수 없습니다.")

    k = min(body.top_k, len(chunks))
    results: list[dict] = []

    for item in body.resumes:
        name = item.filename or "unknown"
        text = (item.text or "").strip()

        if len(text) < 40:
            results.append({
                "filename": name,
                "display_name": re.sub(r"\.[^.]+$", "", name),
                "email": "",
                "score": 0,
                "verdict": "이력서 텍스트 추출 실패 또는 너무 짧음",
                "summary": "",
                "strengths": [],
                "gaps": ["본문 추출을 확인해 주세요."],
                "matched_reference_points": [],
                "rag": {"chunks": [], "chunk_indices": []},
            })
            continue

        # TF-IDF RAG
        picked = pick_top_chunks_tfidf(chunks, text, k)
        excerpts = [p[1][:900] for p in picked]
        indices  = [p[0] for p in picked]

        # Claude 스코어링
        parsed = llm_score_one(client, excerpts, text, name)
        score = max(0, min(100, int(parsed.get("score", 0) or 0)))

        # 이력서 본문에서 이메일 주소 추출
        email_matches = re.findall(
            r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
            text
        )
        # 일반적인 도메인 필터 (이력서 내 실제 이메일만)
        candidate_email = ""
        for em in email_matches:
            # 파일명처럼 생긴 것 제외 (.png .jpg .pdf 등)
            if not re.search(r"\.(png|jpg|jpeg|gif|pdf|doc|xlsx|zip)$", em, re.I):
                candidate_email = em
                break

        results.append({
            "filename": name,
            "display_name": re.sub(r"\.[^.]+$", "", name),
            "email": candidate_email,
            "score": score,
            "verdict": str(parsed.get("verdict", ""))[:300],
            "summary": str(parsed.get("summary", ""))[:1200],
            "strengths": list(parsed.get("strengths") or [])[:5],
            "gaps": list(parsed.get("gaps") or [])[:5],
            "matched_reference_points": list(parsed.get("matched_reference_points") or [])[:5],
            "rag": {
                "chunk_indices": indices,
                "chunks": [
                    {"index": i, "preview": chunks[i][:320], "similarity": round(sim, 4)}
                    for i, _, sim in picked
                ],
            },
        })

    results.sort(key=lambda r: r["score"], reverse=True)
    return {
        "meta": {
            "jd_chars":    len(jd),
            "chunk_count": len(chunks),
            "embed_model": "TF-IDF (sklearn, local)",
            "chat_model":  CHAT_MODEL,
            "top_k":       k,
        },
        "results": results,
    }


# ── HR X: Feedback Vocabulary ─────────────────────────────────

class HRXFeedbackRequest(BaseModel):
    emp_type: str = Field(..., description="직원 유형 (예: 고성과자)")
    scenario: str = Field(..., description="코칭 상황 (예: 강점 칭찬)")
    tone: str = Field(..., description="커뮤니케이션 톤 (예: 코칭형)")
    context: str = Field(default="", description="추가 맥락 (선택)")

@app.post("/api/hrx/feedback")
async def hrx_feedback(req: HRXFeedbackRequest):
    """HR X — Claude AI Feedback Vocabulary Pool 생성"""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=503, detail="ANTHROPIC_API_KEY not configured")

    prompt = f"""당신은 15년 경력의 조직개발 및 HR 코치 전문가입니다.

아래 조건에 맞는 코칭 Vocabulary Pool을 매우 풍부하게 생성해 주세요.

- 직원 유형: {req.emp_type}
- 코칭 상황: {req.scenario}
- 커뮤니케이션 톤: {req.tone}
- 추가 맥락: {req.context or '(없음)'}

다음 7개 섹션을 각각 충분히 풍부하게 작성해 주세요 (각 항목은 불릿으로 구분):

## 1. 핵심 코칭 표현 (이 유형 및 상황에서 가장 효과적인 표현 12개 이상)
## 2. 절대 금기 표현 (써선 안 되는 표현 7개 이상 — 이유와 함께)
## 3. 대화 오프닝 (다양한 방식으로 시작하는 오프닝 6개 이상)
## 4. 완성형 피드백 문장 예시 (구체적 맥락이 담긴 완전한 문장 8개 이상)
## 5. 성찰 유도 코칭 질문 (열린 질문 형식 10개 이상)
## 6. 마무리 표현 (동기부여 및 다음 행동 연결 6개 이상)
## 7. 감정 인식 및 공감 표현 (7개 이상)

모든 표현은 {req.emp_type} 유형의 특성을 깊이 반영해 주세요."""

    try:
        client = _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model=CHAT_MODEL,
            max_tokens=3500,
            messages=[{"role": "user", "content": prompt}],
        )
        return {"result": msg.content[0].text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ── HR X: 직접 작성 모드 — 자유 상황 텍스트 → Claude 분석 ─────
class FeedbackDirectRequest(BaseModel):
    situation_text: str

@app.post("/api/hrx/feedback-direct")
async def hrx_feedback_direct(req: FeedbackDirectRequest):
    """직접 작성 상황 텍스트를 Claude AI가 분석하여 맞춤 Vocabulary 생성"""
    if not req.situation_text or len(req.situation_text.strip()) < 20:
        raise HTTPException(status_code=400, detail="상황 텍스트가 너무 짧습니다.")

    prompt = f"""당신은 한화투자증권 혁신(리더)지원실 소속 HR 코칭 전문가입니다. 아래 코칭 상황을 깊이 이해하고, 이 상황에 최적화된 피드백 Vocabulary Pool을 생성해 주세요.

[코칭 상황]
{req.situation_text}

위 상황을 분석하여 다음 7개 섹션을 한국어로 작성해 주세요. 각 항목은 불릿(·)으로 구분하고, 이 상황의 맥락과 뉘앙스를 충분히 반영해 주세요.

중요: 응답의 첫 줄에 전체 제목(예: "# 한국 기업 HR 코칭 Vocabulary Pool" 등)을 절대 추가하지 마세요. 아래 섹션 헤더(##)만 사용하세요.

## 1. 상황 분석 요약 (이 상황의 핵심 코칭 포인트 3~4가지)
## 2. 핵심 코칭 표현 (이 상황에서 가장 효과적인 표현 10개 이상)
## 3. 절대 금기 표현 (이 상황에서 써선 안 되는 표현 5개 이상 — 이유 포함)
## 4. 대화 오프닝 스크립트 (이 상황에 맞는 시작 문장 5개 이상)
## 5. 완성형 피드백 문장 예시 (이 맥락이 담긴 완전한 문장 7개 이상)
## 6. 성찰 유도 코칭 질문 (이 상황에 맞는 열린 질문 8개 이상)
## 7. 마무리 · 동기부여 표현 (다음 행동으로 연결하는 마무리 5개 이상)

모든 내용은 위 상황의 구체적 맥락을 반영하여 실제로 사용 가능한 표현으로 작성해 주세요."""

    try:
        client = _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model=CHAT_MODEL,
            max_tokens=3500,
            messages=[{"role": "user", "content": prompt}],
        )
        raw_text = msg.content[0].text

        # 첫 줄이 # 로 시작하는 전체 제목이면 제거 (## 섹션 헤더는 유지)
        lines = raw_text.split('\n')
        filtered_lines = []
        for i, line in enumerate(lines):
            stripped = line.strip()
            # 첫 번째 실질 줄이 # 제목(##이 아닌 단일 #)이면 건너뜀
            if i == 0 and stripped.startswith('#') and not stripped.startswith('##'):
                continue
            filtered_lines.append(line)
        result_text = '\n'.join(filtered_lines).lstrip('\n')

        return {"result": result_text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ── HR X: 프로젝트 스킬 분류 ─────────────────────────────────
class ClassifyProjectRequest(BaseModel):
    name: str
    desc: str
    skills_raw: str

@app.post("/api/hrx/classify-project")
async def hrx_classify_project(req: ClassifyProjectRequest):
    """프로젝트 설명 + 자유 스킬 텍스트 → 필수/우대 분류 + 기간/팀 추천"""
    prompt = f"""당신은 IT 프로젝트 기술 스펙 분석 전문가입니다.

프로젝트명: {req.name}
프로젝트 설명: {req.desc}
입력된 기술/요구사항 (자유 형식): {req.skills_raw}

위 정보를 분석하여 다음 JSON만 출력하세요 (다른 텍스트 절대 없음):
{{
  "required": ["필수스킬1", "필수스킬2", ...],
  "preferred": ["우대스킬1", "우대스킬2", ...],
  "suggested_months": 숫자(1~18),
  "suggested_team": 숫자(1~10),
  "reasoning": "한 줄 추천 근거"
}}

기준:
- required: 프로젝트 핵심 기능 구현에 반드시 필요한 기술
- preferred: 있으면 좋지만 없어도 되는 기술
- suggested_months: 프로젝트 규모와 복잡도 기반 권장 기간
- suggested_team: 필요 기술 다양성과 규모 기반 권장 인원
"""
    try:
        client = _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model=CHAT_MODEL,
            max_tokens=800,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text
        json_match = re.search(r"\{[\s\S]*\}", raw)
        if json_match:
            return json.loads(json_match.group(0))
        return {"required": [], "preferred": [], "suggested_months": 4, "suggested_team": 4, "reasoning": ""}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ── HR X: Word(.docx) 다운로드 ────────────────────────────────
class DocxDownloadRequest(BaseModel):
    title: str = "Feedback Vocabulary"
    subtitle: str = ""
    content: str  # 평문 텍스트 (섹션 구분은 ## 헤더로)

@app.post("/api/hrx/download-docx")
async def hrx_download_docx(req: DocxDownloadRequest):
    """Vocabulary 결과를 Word(.docx)로 변환하여 반환"""
    try:
        from docx import Document as _Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io as _io

        doc = _Document()

        # 기본 여백 설정
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        # 제목
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(req.title)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0xEA, 0x58, 0x0C)  # orange

        # 부제목 (상황 요약)
        if req.subtitle:
            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_para.add_run(f"[상황] {req.subtitle}{'...' if len(req.subtitle) >= 120 else ''}")
            sub_run.font.size = Pt(10)
            sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_paragraph()  # 공백

        # 본문 파싱
        lines = req.content.split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue

            # ## 섹션 헤더
            if stripped.startswith('##'):
                heading_text = stripped.lstrip('#').strip()
                h_para = doc.add_paragraph()
                h_run = h_para.add_run(heading_text)
                h_run.bold = True
                h_run.font.size = Pt(12)
                h_run.font.color.rgb = RGBColor(0xEA, 0x58, 0x0C)
                # 헤더 위 공백
                h_para.paragraph_format.space_before = Pt(12)
                h_para.paragraph_format.space_after = Pt(4)

            # 불릿 항목 (·, -, •, *)
            elif stripped[0] in ('·', '-', '•', '*') and len(stripped) > 2:
                item_text = stripped[1:].strip().lstrip('-').strip()
                item_para = doc.add_paragraph(style='List Bullet')
                item_run = item_para.add_run(item_text)
                item_run.font.size = Pt(10.5)

            # 숫자 리스트
            elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in '.):':
                item_text = stripped[2:].strip()
                item_para = doc.add_paragraph(style='List Number')
                item_run = item_para.add_run(item_text)
                item_run.font.size = Pt(10.5)

            # 일반 텍스트
            else:
                p = doc.add_paragraph()
                p.add_run(stripped).font.size = Pt(10.5)

        # 생성 날짜 footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer_para.add_run(f"Generated by HR Room · {__import__('datetime').date.today()}")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)

        # 메모리에 저장
        buf = _io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        from fastapi.responses import StreamingResponse
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="FeedbackVocabulary.docx"'}
        )
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx 패키지가 설치되어 있지 않습니다.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ── HR X: Analytics (연봉 리뷰) ──────────────────────────────

_HRX_EXCEL = ROOT / "data" / "hr_dummy_salary_review_completed_10types.xlsx"

_POSITIVE_KW = ["훌륭","우수","탁월","기여","성과","완료","달성","향상","개선",
                "리더십","주도","성장","발전","협력","소통","신뢰","책임","창의","혁신","효율",
                "주도적","적극적","선도","완성도","안정적"]
_NEGATIVE_KW = ["미흡","부족","지연","이슈","문제","아쉬움","부진","실수",
                "오류","지체","혼선","갈등","불명확","소통부재","방어적","수동적",
                "미달","불충분","재작업","위험"]

def _hrx_sentiment(text: str) -> float:
    if not isinstance(text, str) or text.strip() in ("", "nan", "None"):
        return 0.0
    pos = sum(1 for w in _POSITIVE_KW if w in text)
    neg = sum(1 for w in _NEGATIVE_KW if w in text)
    total = pos + neg
    return round((pos - neg) / total * 100, 1) if total else 0.0

def _hrx_specificity(text: str) -> float:
    import re as _re
    if not isinstance(text, str): return 0.0
    nums = len(_re.findall(r"\d+", text))
    pcts = len(_re.findall(r"\d+%", text))
    return min(100.0, nums * 5 + pcts * 10 + len(text) / 12)

@app.get("/api/hrx/analytics")
def hrx_analytics():
    """HR X — 실제 Excel 데이터 기반 연봉 Analytics"""
    try:
        import pandas as _pd
        df = _pd.read_excel(_HRX_EXCEL, sheet_name="Sheet1")
        rename = {
            "구분": "team", "일련번호": "employee_id", "직급": "grade",
            "26년 연봉": "salary_2026", "25년 연봉": "salary_2025", "24년 연봉": "salary_2024",
            "26년 2분기 피드백": "fb_q2", "26년 3분기 피드백": "fb_q3", "26년 4분기 피드백": "fb_q4",
        }
        df = df.rename(columns=rename)
        df["raise_26"] = (df["salary_2026"] - df["salary_2025"]) / df["salary_2025"] * 100
        df["raise_25"] = (df["salary_2025"] - df["salary_2024"]) / df["salary_2024"] * 100
        df["fb_score"] = df.apply(
            lambda r: (_hrx_sentiment(str(r.fb_q2)) + _hrx_sentiment(str(r.fb_q3)) + _hrx_sentiment(str(r.fb_q4))) / 3,
            axis=1,
        )
        df["fb_len"] = df.apply(
            lambda r: (len(str(r.fb_q2)) + len(str(r.fb_q3)) + len(str(r.fb_q4))) / 3,
            axis=1,
        )
        df["fb_spec"] = df.apply(
            lambda r: (_hrx_specificity(str(r.fb_q2)) + _hrx_specificity(str(r.fb_q3)) + _hrx_specificity(str(r.fb_q4))) / 3,
            axis=1,
        )
        # KPIs
        total = len(df)
        avg_salary = float(df["salary_2026"].mean())
        avg_raise = float(df["raise_26"].mean())
        corr = float(df[["fb_score", "raise_26"]].corr().iloc[0, 1])

        # scatter points (all employees)
        scatter = [{"x": round(float(r.fb_score), 1), "y": round(float(r.raise_26), 2),
                    "team": str(r.team), "id": str(r.employee_id), "grade": str(r.grade)}
                   for _, r in df.iterrows()]

        # team aggregates
        tg = df.groupby("team").agg(
            count=("employee_id", "count"),
            avg_salary_2024=("salary_2024", "mean"),
            avg_salary_2025=("salary_2025", "mean"),
            avg_salary_2026=("salary_2026", "mean"),
            raise_26=("raise_26", "mean"),
            raise_25=("raise_25", "mean"),
            fb_score=("fb_score", "mean"),
            fb_len=("fb_len", "mean"),
            fb_spec=("fb_spec", "mean"),
        ).reset_index().round(2)
        teams = tg.to_dict(orient="records")

        # per-employee (for individual tab)
        employees = []
        for _, r in df.iterrows():
            employees.append({
                "employee_id": str(r.employee_id), "team": str(r.team), "grade": str(r.grade),
                "salary_2024": float(r.salary_2024), "salary_2025": float(r.salary_2025), "salary_2026": float(r.salary_2026),
                "raise_26": round(float(r.raise_26), 2), "raise_25": round(float(r.raise_25), 2),
                "fb_q2": str(r.fb_q2), "fb_q3": str(r.fb_q3), "fb_q4": str(r.fb_q4),
                "fb_score": round(float(r.fb_score), 1),
                "fb_len": round(float(r.fb_len), 1),
                "fb_spec": round(float(r.fb_spec), 1),
            })

        return {
            "ok": True, "total": total, "avg_salary": avg_salary,
            "avg_raise": avg_raise, "corr": corr,
            "scatter": scatter, "teams": teams, "employees": employees,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class HRXSalaryAIRequest(BaseModel):
    employee_id: str
    team: str
    grade: str
    salary_2024: float
    salary_2025: float
    salary_2026: float
    raise_26: float
    raise_25: float
    fb_q2: str
    fb_q3: str
    fb_q4: str
    fb_score: float
    fb_len: float
    fb_spec: float

@app.post("/api/hrx/salary-ai")
async def hrx_salary_ai(req: HRXSalaryAIRequest):
    """HR X — Claude AI 개인 연봉 분석 (2027년 추천)"""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=503, detail="ANTHROPIC_API_KEY not configured")
    prompt = f"""당신은 보상 및 성과 분석 전문 HR 컨설턴트입니다.

아래 직원 데이터를 분석하고 2027년 연봉 적정 수준을 제안해 주세요.

직원 정보
- ID: {req.employee_id} / 팀: {req.team} / 직급: {req.grade}
- 2024년 연봉: {req.salary_2024:,.0f}원
- 2025년 연봉: {req.salary_2025:,.0f}원 (인상률 {req.raise_25:.1f}%)
- 2026년 연봉: {req.salary_2026:,.0f}원 (인상률 {req.raise_26:.1f}%)

분기별 피드백 원문
- Q2: {req.fb_q2}
- Q3: {req.fb_q3}
- Q4: {req.fb_q4}

자동 분석 값
- 종합 감정 점수: {req.fb_score:.1f} / 피드백 평균 길이: {req.fb_len:.0f}자 / 구체성: {req.fb_spec:.1f}

다음 5개 섹션으로 분석해 주세요:

## 1. 피드백 감정 심층 분석 (각 분기별 핵심 키워드 및 전반적 평가)
## 2. 성과-보상 정합성 평가 (현재 인상률이 피드백 수준과 맞는지)
## 3. 2027년 연봉 추천 (구체적 금액 범위와 인상률 제시)
## 4. 추천 근거 (3가지 이상 명확한 논거)
## 5. 리스크 및 권고사항 (이탈 위험, 동기부여 포인트 등)

모든 분석은 한국어로 작성해 주세요."""
    try:
        client = _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model=CHAT_MODEL, max_tokens=1800,
            messages=[{"role": "user", "content": prompt}],
        )
        return {"result": msg.content[0].text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ══════════════════════════════════════════════════════════════
# 합격자 이메일/카카오톡 발송 API
# ══════════════════════════════════════════════════════════════

class NotificationRequest(BaseModel):
    type: str  # "email"
    candidates: list[str]           # 합격자 이름 목록
    recipients: list[str] = []      # 실제 수신자 이메일 주소 목록
    from_email: str = ""            # 발신 Gmail 주소
    smtp_pass_input: str = ""       # Gmail 앱 비밀번호 (UI에서 전달)
    subject: str = ""
    body: str = ""

@app.post("/api/send-notification")
async def send_notification(req: NotificationRequest):
    """합격자 이메일 발송 (Gmail SMTP)"""
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart

    if req.type != "email":
        return {"ok": False, "error": "지원하지 않는 발송 유형"}

    # SMTP 설정: 환경변수 우선, 없으면 UI 입력값 사용
    smtp_host = os.getenv("SMTP_HOST", "smtp.gmail.com")
    smtp_port = int(os.getenv("SMTP_PORT", "587"))
    smtp_user = os.getenv("SMTP_USER", "") or req.from_email
    smtp_pass = os.getenv("SMTP_PASS", "") or req.smtp_pass_input

    if not smtp_pass:
        return {"ok": False, "error": "Gmail 앱 비밀번호가 입력되지 않았습니다."}
    if not smtp_user:
        return {"ok": False, "error": "발신 Gmail 주소가 입력되지 않았습니다."}
    if not req.recipients:
        return {"ok": False, "error": "수신자 이메일이 없습니다."}

    sent, failed = [], []
    for to_addr in req.recipients:
        try:
            msg = MIMEMultipart("alternative")
            msg["Subject"] = req.subject or f"[HR Room] 채용 합격 안내"
            msg["From"]    = smtp_user
            msg["To"]      = to_addr
            msg.attach(MIMEText(req.body or "", "plain", "utf-8"))

            with smtplib.SMTP(smtp_host, smtp_port) as s:
                s.starttls()
                s.login(smtp_user, smtp_pass)
                s.sendmail(smtp_user, [to_addr], msg.as_string())
            sent.append(to_addr)
        except Exception as ex:
            failed.append({"email": to_addr, "error": str(ex)})

    return {"ok": True, "sent": sent, "failed": failed}


# ==============================================
# 헬스체크
# ==============================================
@app.get("/api/health")
def health():
    return {
        "ok": bool(ANTHROPIC_API_KEY),
        "openai_configured": bool(ANTHROPIC_API_KEY),
        "anthropic_configured": bool(ANTHROPIC_API_KEY),
        "model": CHAT_MODEL,
        "embed": "TF-IDF (local)",
    }


# ==============================================
# 정적 파일 서빙 (모든 API 라우트 후 마지막에)
# ==============================================
_TUITION_STATIC = ROOT / "static" / "tuition"
if _TUITION_STATIC.exists():
    app.mount("/tuition", StaticFiles(directory=str(_TUITION_STATIC)), name="tuition")

app.mount("/", StaticFiles(directory=str(ROOT / "static"), html=True), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("server:app", host="127.0.0.1", port=8765, reload=True)
